# # ai_developer_agent.py
# import google.generativeai as genai
# import json
# import time

# class AIDeveloperAgent:
#     def __init__(self, model_name="gemini-2.5-flash"):
#         print("[INIT] Initializing AI Developer Agent...")
#         self.model = genai.GenerativeModel(model_name)
#         print(f"✅ AI Developer Agent initialized using model: {model_name}")

#     def analyze_code(self, file_name, file_content):
#         """Stage 1 & 2: Perception + Reasoning"""
#         print(f"\n[AI] Analyzing file: {file_name}")
#         prompt = f"""
#         You are an AI code review assistant.
#         Analyze the following Python file and return structured JSON with:
#         1. Issues (like unused imports, bad naming, redundancy)
#         2. Refactor Suggestions
#         3. Documentation Summary

#         Format your response strictly in JSON:
#         {{
#             "file": "{file_name}",
#             "issues": [
#                 {{
#                     "line": <line_number>,
#                     "type": "<issue_type>",
#                     "suggestion": "<suggestion>"
#                 }}
#             ],
#             "refactors": [
#                 {{
#                     "description": "<refactor_description>",
#                     "impact": "<benefit>"
#                 }}
#             ],
#             "summary": "<short summary>"
#         }}

#         Code:
#         {file_content}
#         """
#         try:
#             response = self.model.generate_content(prompt)
#             text = response.text.strip()
#             json_data = json.loads(text)
#             print(f"✅ Analysis complete for {file_name}")
#             return json_data
#         except Exception as e:
#             print(f"⚠️ Error analyzing {file_name}: {e}")
#             return {
#                 "file": file_name,
#                 "issues": [],
#                 "refactors": [],
#                 "summary": f"Analysis failed: {e}"
#             }

#     def apply_suggestions(self, file_content, accepted_suggestions):
#         """Stage 3: Action — apply accepted AI changes"""
#         print("[AI] Applying accepted suggestions...")
#         # For now, simulate transformation (later can use AST or LLM code edit)
#         for suggestion in accepted_suggestions:
#             file_content += f"\n# Applied change: {suggestion['suggestion']}"
#         print("✅ Applied accepted suggestions.")
#         return file_content
    
#     def run_full_pipeline(self, repo_url, accepted_suggestions=None):
#         """
#         Complete AI Developer flow:
#         1. Fetch repo
#         2. Analyze code
#         3. Apply accepted suggestions (if any)
#         4. Generate documentation
#         """
#         from main import fetch_repo  # import existing route logic dynamically
#         from documentation_agent import DocumentationAgent

#         print(f"\n[PIPELINE] Running full developer agent pipeline for repo: {repo_url}")
#         docs_agent = DocumentationAgent()

#         # 1️⃣ Fetch repo
#         repo_data = fetch_repo(repo_url)
#         files = repo_data["files"]

#         # 2️⃣ Analyze code
#         all_analysis = []
#         for f in files:
#             analysis = self.analyze_code(f["file_name"], f["content"])
#             all_analysis.append(analysis)

#         # 3️⃣ Apply suggestions (simulate)
#         if accepted_suggestions:
#             print("[ACTION] Applying accepted suggestions...")
#             for s in accepted_suggestions:
#                 fmatch = next((f for f in files if f["file_name"] == s["file_name"]), None)
#                 if fmatch:
#                     fmatch["content"] = self.apply_suggestions(fmatch["content"], [s])
#         else:
#             print("ℹ️ No accepted suggestions provided. Skipping modification.")

#         # 4️⃣ Generate documentation
#         docs_output = []
#         for f in files:
#             doc = docs_agent.generate_docs_for_file(f["file_name"], f["content"])
#             docs_output.append({
#                 "file_name": f["file_name"],
#                 "documentation": doc
#             })

#         return {
#             "repo_name": repo_data["repo_name"],
#             "analysis": all_analysis,
#             "documentation": docs_output
#         }

# # developer_agent.py
# import os, json, time, requests
# from autogen import AssistantAgent, register_function
# import google.generativeai as genai
# from documentation_agent import DocumentationAgent

# # -------------------------
# # Initialize Gemini API
# # -------------------------
# genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# # -------------------------
# # Tool: Fetch Repo
# # -------------------------
# @register_function
# def fetch_repo(repo_url: str):
#     """Fetch all code files from a GitHub repo and return structured data."""
#     print(f"\n[FETCH] Fetching repo: {repo_url}")
#     parts = repo_url.strip("/").split("/")
#     username, reponame = parts[-2], parts[-1]
#     api_url = f"https://api.github.com/repos/{username}/{reponame}/contents"

#     response = requests.get(api_url)
#     response.raise_for_status()
#     repo_data = response.json()
#     allowed_ext = [".py", ".js", ".java", ".ts", ".cpp", ".c", ".html", ".css", ".json", ".xml", ".md"]

#     files = []
#     for f in repo_data:
#         if f["type"] == "file" and any(f["name"].endswith(ext) for ext in allowed_ext):
#             content = requests.get(f["download_url"]).text
#             files.append({"file_name": f["name"], "content": content})
#     print(f"✅ Repo fetched successfully: {len(files)} files")
#     return {"repo_name": reponame, "files": files}


# # -------------------------
# # Tool: Commit Changes (Pseudo)
# # -------------------------
# @register_function
# def commit_to_github(repo_url: str, file_name: str, new_content: str, comment: str):
#     """Simulate committing code to GitHub (can extend to real API later)."""
#     print(f"[COMMIT] {file_name} → Commit message: {comment}")
#     return f"Committed {file_name} successfully to {repo_url}"


# # -------------------------
# # Tool: Generate Documentation
# # -------------------------
# @register_function
# def generate_docs(file_name: str, file_content: str):
#     """Generate documentation for a single file."""
#     doc_agent = DocumentationAgent()
#     doc = doc_agent.generate_docs_for_file(file_name, file_content)
#     return doc


# # -------------------------
# # Define Developer Agent
# # -------------------------
# class AIDeveloperAgent:
#     def __init__(self, model_name="gemini-2.5-flash"):
#         print("[INIT] Initializing Autogen Developer Agent...")
#         self.model_name = model_name

#         self.agent = AssistantAgent(
#             name="DeveloperAgent",
#             model=model_name,
#             system_message=(
#                 "You are an expert AI software developer. "
#                 "Your goal is to: "
#                 "1. Fetch the given GitHub repo "
#                 "2. Analyze the code for issues/refactors "
#                 "3. Apply accepted suggestions "
#                 "4. Generate documentation for all files "
#                 "5. Return a complete structured JSON response"
#             ),
#             tools=[fetch_repo, commit_to_github, generate_docs],
#         )

#     def run_pipeline(self, repo_url: str, accepted_suggestions=None):
#         """Execute full developer flow autonomously via Autogen."""
#         print(f"\n🚀 [RUN] DeveloperAgent pipeline started for: {repo_url}")
#         start = time.time()

#         # Construct task for Autogen agent
#         prompt = f"""
#         You are the DeveloperAgent responsible for automating the software analysis workflow.
#         Step 1: Use fetch_repo('{repo_url}') to get all files.
#         Step 2: For each file, analyze and find issues or optimizations.
#         Step 3: If accepted suggestions are provided, apply them and call commit_to_github().
#         Step 4: Finally, call generate_docs() for each file and return:
#         {{
#             "repo_name": "<repo_name>",
#             "analysis": [...],
#             "documentation": [...]
#         }}
#         """

#         # Add accepted suggestions context (if any)
#         if accepted_suggestions:
#             prompt += f"\nAccepted Suggestions:\n{json.dumps(accepted_suggestions, indent=2)}"

#         result = self.agent.run(task=prompt)
#         elapsed = round(time.time() - start, 2)
#         print(f"✅ DeveloperAgent pipeline completed in {elapsed}s")
#         return result

# # developer_agent.py
# import json, time, requests, tempfile, os
# import google.generativeai as genai
# from autogen import ConversableAgent
# from fastapi.responses import FileResponse


# class DeveloperAgent(ConversableAgent):
#     def __init__(self, model_name="gemini-2.5-flash"):
#         super().__init__(name="DeveloperAgent", description="Autonomous AI Developer Agent for CodeZen.")
#         self.model_name = model_name
#         self.model = genai.GenerativeModel(model_name)
#         print(f"✅ DeveloperAgent initialized with {model_name}")

#     # ----------------------------------------
#     # 1️⃣ Repo Fetching (Recursive)
#     # ----------------------------------------
#     def fetch_repo(self, repo_url: str):
#         print(f"[FETCH] Fetching repository: {repo_url}")
#         parts = repo_url.strip("/").split("/")
#         username, reponame = parts[-2], parts[-1].replace(".git", "")
#         api_url = f"https://api.github.com/repos/{username}/{reponame}/contents"

#         def fetch_directory(api_url):
#             response = requests.get(api_url)
#             if response.status_code != 200:
#                 raise Exception("Failed to fetch repo contents")
#             items = response.json()
#             files = []
#             allowed_ext = [".py", ".js", ".java", ".ts", ".cpp", ".c", ".html", ".css"]
#             for item in items:
#                 if item["type"] == "file" and any(item["name"].endswith(ext) for ext in allowed_ext):
#                     content = requests.get(item["download_url"]).text
#                     files.append({"file_name": item["path"], "content": content})
#                 elif item["type"] == "dir":
#                     files.extend(fetch_directory(item["url"]))
#             return files

#         all_files = fetch_directory(api_url)
#         print(f"✅ Repo fetched ({len(all_files)} code files).")
#         return {"repo_name": reponame, "files": all_files}

#     # ----------------------------------------
#     # 2️⃣ Code Analysis
#     # ----------------------------------------
#     def analyze_code(self, file_name, file_content):
#         print(f"[AI] Analyzing {file_name}")
#         prompt = f"""
#         You are an AI code reviewer.
#         Analyze the code and return JSON with:
#         - issues
#         - refactor suggestions
#         - summary

#         Code:
#         {file_content}
#         """
#         try:
#             res = self.model.generate_content(prompt)
#             return json.loads(res.text)
#         except Exception as e:
#             print(f"⚠️ Analysis failed: {e}")
#             return {"file": file_name, "issues": [], "refactors": [], "summary": str(e)}

#     # ----------------------------------------
#     # 3️⃣ Apply Suggestions
#     # ----------------------------------------
#     def apply_suggestions(self, file_content, suggestions):
#         print("[ACTION] Applying accepted suggestions...")
#         for s in suggestions:
#             file_content += f"\n# Applied suggestion: {s.get('comment','')}"
#         return file_content

#     # ----------------------------------------
#     # 4️⃣ Documentation Generation
#     # ----------------------------------------
#     def generate_docs(self, file_name, file_content):
#         print(f"[DOCS] Generating documentation for {file_name}")
#         prompt = f"""
#         Generate documentation for this file in markdown format.
#         Include:
#         - Purpose
#         - Key Functions
#         - Example Usage
#         - Summary

#         Code:
#         {file_content}
#         """
#         try:
#             res = self.model.generate_content(prompt)
#             return res.text.strip()
#         except Exception as e:
#             return f"Error generating docs: {e}"

#     # ----------------------------------------
#     # 5️⃣ Export Documentation
#     # ----------------------------------------
#     def export_docs(self, docs_output, format="md"):
#         print(f"[EXPORT] Exporting docs as {format}")
#         if format.lower() == "md":
#             path = tempfile.mktemp(suffix=".md")
#             with open(path, "w", encoding="utf-8") as f:
#                 for doc in docs_output:
#                     f.write(f"# {doc['file_name']}\n\n{doc['documentation']}\n\n---\n")
#             print(f"📄 Exported docs at: {path}")
#             return path  # Return path for reference instead of FileResponse

#     # ----------------------------------------
#     # 🚀 Full AutoGen-Orchestrated Pipeline
#     # ----------------------------------------
#     def run_full_pipeline(self, repo_url, accepted_suggestions=None):
#         print(f"\n🚀 Running full AutoGen pipeline for {repo_url}")
#         repo_data = self.fetch_repo(repo_url)
#         files = repo_data["files"]

#         all_analysis, docs_output = [], []

#         for f in files:
#             analysis = self.analyze_code(f["file_name"], f["content"])
#             all_analysis.append(analysis)

#             if accepted_suggestions:
#                 matched = [s for s in accepted_suggestions if s["file_name"] == f["file_name"]]
#                 if matched:
#                     f["content"] = self.apply_suggestions(f["content"], matched)

#             doc_text = self.generate_docs(f["file_name"], f["content"])
#             docs_output.append({"file_name": f["file_name"], "documentation": doc_text})

#         # 🧾 Export documentation to Markdown
#         export_path = self.export_docs(docs_output, format="md")

#         print("✅ Pipeline completed successfully.")
#         return {
#             "repo": repo_data["repo_name"],
#             "analysis": all_analysis,
#             "documentation": docs_output,
#             "exported_file": export_path,
#         }


# # developer_agent.py
# import json, time, requests, os, re
# import google.generativeai as genai
# from autogen import ConversableAgent


# class DeveloperAgent(ConversableAgent):
#     def __init__(self, model_name="gemini-2.5-flash"):
#         super().__init__(name="DeveloperAgent", description="Unified AI Developer Agent for CodeZen.")
#         self.model_name = model_name
#         self.model = genai.GenerativeModel(model_name)
#         self.user_preferences = {}  # 🧠 Track accept/reject patterns
#         print(f"✅ DeveloperAgent initialized with {model_name}")

#     # ----------------------------------------
#     # 1️⃣ PERCEPTION STAGE — Repository Fetching (Full Tree)
#     # ----------------------------------------
#     def fetch_repo(self, repo_url: str):
#         print(f"\n[PERCEPTION] 🌐 Fetching repository structure from: {repo_url}")
#         parts = repo_url.strip("/").split("/")
#         username, reponame = parts[-2], parts[-1].replace(".git", "")

#         # Detect default branch automatically
#         repo_meta_url = f"https://api.github.com/repos/{username}/{reponame}"
#         meta_resp = requests.get(repo_meta_url)
#         if meta_resp.status_code != 200:
#             raise Exception(f"❌ [PERCEPTION] Failed to fetch repo metadata: {meta_resp.status_code}")
#         default_branch = meta_resp.json().get("default_branch", "main")
#         print(f"[PERCEPTION] 🪶 Default branch detected: {default_branch}")

#         # Use Git Tree API for recursive file fetch
#         tree_api_url = f"https://api.github.com/repos/{username}/{reponame}/git/trees/{default_branch}?recursive=1"
#         tree_resp = requests.get(tree_api_url)
#         if tree_resp.status_code != 200:
#             raise Exception(f"❌ [PERCEPTION] Failed to fetch repo tree: {tree_resp.status_code}")

#         tree_data = tree_resp.json()
#         files = []
#         allowed_ext = [".py", ".js", ".java", ".ts", ".cpp", ".c", ".html", ".css", ".ipynb"]

#         print(f"[PERCEPTION] 🔍 Scanning {len(tree_data.get('tree', []))} items from repo tree...")
#         for item in tree_data.get("tree", []):
#             if item["type"] == "blob" and any(item["path"].endswith(ext) for ext in allowed_ext):
#                 raw_url = f"https://raw.githubusercontent.com/{username}/{reponame}/{default_branch}/{item['path']}"
#                 try:
#                     content = requests.get(raw_url).text

#                     # ✅ Handle Jupyter Notebooks (.ipynb)
#                     if item["path"].endswith(".ipynb"):
#                         try:
#                             nb_data = json.loads(content)
#                             code_cells = [
#                                 "\n".join(cell.get("source", []))
#                                 for cell in nb_data.get("cells", [])
#                                 if cell.get("cell_type") == "code"
#                             ]
#                             content = "\n\n".join(code_cells)
#                             print(f"📘 [NOTEBOOK] Extracted {len(code_cells)} code cells from {item['path']}")
#                         except Exception as e:
#                             print(f"⚠️ [NOTEBOOK] Failed to parse notebook {item['path']}: {e}")

#                     files.append({"file_name": item["path"], "content": content})

#                 except Exception as e:
#                     print(f"⚠️ [PERCEPTION] Failed to fetch {item['path']}: {e}")

#         print(f"✅ [PERCEPTION] Repository fetched successfully — {len(files)} code files detected.")
#         return {"repo_name": reponame, "files": files}

#     # ----------------------------------------
#     # 2️⃣ REASONING STAGE — Code Analysis with JSON Recovery
#     # ----------------------------------------
#     def analyze_code(self, file_name, file_content):
#         print(f"\n[REASONING] 🧩 Analyzing file: {file_name}")

#         prompt = f"""
#         You are an expert AI code reviewer.
#         Carefully analyze the following code and respond STRICTLY in valid JSON format.

#         JSON Schema:
#         {{
#             "issues": ["list of potential issues or errors found"],
#             "refactors": ["list of suggested improvements"],
#             "summary": "short summary of what this file does"
#         }}

#         Do not include markdown or explanations outside JSON.

#         Code:
#         {file_content}
#         """

#         try:
#             res = self.model.generate_content(prompt)
#             raw_output = res.text.strip()

#             try:
#                 parsed = json.loads(raw_output)
#                 print(f"✅ [REASONING] Parsed valid JSON for {file_name}")
#             except json.JSONDecodeError:
#                 print(f"⚠️ [REASONING] Model returned non-JSON output. Attempting recovery...")
#                 match = re.search(r"\{.*\}", raw_output, re.DOTALL)
#                 if match:
#                     try:
#                         parsed = json.loads(match.group())
#                         print("✅ [REASONING] Successfully extracted JSON structure.")
#                     except Exception:
#                         parsed = {"issues": [], "refactors": [], "summary": raw_output[:300]}
#                         print("⚠️ [REASONING] Fallback: partial summary used.")
#                 else:
#                     parsed = {"issues": [], "refactors": [], "summary": raw_output[:300]}
#                     print("⚠️ [REASONING] No JSON found; fallback summary used.")

#             print(f"🔍 [DEBUG] Issues: {len(parsed.get('issues', []))}, Refactors: {len(parsed.get('refactors', []))}")
#             return parsed

#         except Exception as e:
#             print(f"❌ [REASONING] Failed to analyze {file_name}: {e}")
#             return {"file": file_name, "issues": [], "refactors": [], "summary": str(e)}

#     # ----------------------------------------
#     # 3️⃣ ACTION STAGE — Apply Suggestions + (Simulated) Commit
#     # ----------------------------------------
#     def apply_suggestions(self, file_content, suggestions):
#         print(f"\n[ACTION] 🛠️ Applying accepted suggestions ({len(suggestions)} items)...")
#         for s in suggestions:
#             file_content += f"\n# Applied suggestion: {s.get('comment','')}"

#         print("[ACTION] ✅ Suggestions applied.")
#         return file_content

#     def commit_to_github(self, repo_url, file_path, updated_content, commit_message):
#         print(f"[ACTION] 🪶 Simulating commit to GitHub for {file_path} ...")
#         time.sleep(1)
#         print(f"[ACTION] ✅ Simulated commit complete for {file_path}")
#         return True

#     # ----------------------------------------
#     # 4️⃣ DOCUMENTATION STAGE — File-Level Documentation
#     # ----------------------------------------
#     def generate_docs(self, file_name, file_content, format="md"):
#         print(f"\n[DOCUMENTATION] 🧾 Generating {format.upper()} documentation for {file_name}")
#         doc_prompt = f"""
#         You are a senior software architect.
#         Generate a comprehensive Technical Design Document (TTD) for this file in {format.upper()} format.
#         Include:
#         - Purpose of this file
#         - Key functions and classes
#         - Data flow or dependencies
#         - Example usage
#         - Design summary

#         Code:
#         {file_content}
#         """
#         try:
#             res = self.model.generate_content(doc_prompt)
#             print(f"[DOCUMENTATION] ✅ Documentation generated successfully for {file_name}")
#             return res.text.strip()
#         except Exception as e:
#             print(f"❌ [DOCUMENTATION] Error generating docs for {file_name}: {e}")
#             return f"Error generating docs: {e}"

#     # ----------------------------------------
#     # 🧱 4B️⃣ PROJECT-LEVEL DOCUMENTATION GENERATOR
#     # ----------------------------------------
#     def generate_project_docs(self, repo_name, all_files, analyses, format="md"):
#         print(f"\n[PROJECT DOCS] 🏗️ Generating overall {format.upper()} documentation for project: {repo_name}")

#         combined_code = ""
#         for f in all_files:
#             combined_code += f"\n\n### FILE: {f['file_name']}\n{f['content']}\n"

#         summary_context = "\n\n".join(
#             [f"File: {a.get('file', 'N/A')} — {a.get('summary', '')}" for a in analyses]
#         )

#         project_prompt = f"""
#         You are a senior software architect and system analyst.
#         You are given a full project repository named **{repo_name}**.
#         Below are all code files, scripts, and their analyses.

#         Your task:
#         Write a **comprehensive project-level Technical Design Document (TTD)** that explains:
#         - Project overview and purpose
#         - System architecture and major components
#         - How modules or files interact
#         - Data flow, control flow, and dependencies
#         - Technologies and libraries used
#         - Key functions, APIs, or classes
#         - Execution flow (end-to-end)
#         - Example usage (how it runs or what it does)
#         - Potential improvements or limitations

#         Use a structured {format.upper()} format.
#         Avoid per-file repetition; focus on the holistic system.

#         ANALYSES:
#         {summary_context}

#         CODEBASE:
#         {combined_code}
#         """
#         try:
#             res = self.model.generate_content(project_prompt)
#             print(f"[PROJECT DOCS] ✅ Project-level documentation generated successfully.")
#             return res.text.strip()
#         except Exception as e:
#             print(f"❌ [PROJECT DOCS] Error generating project-level docs: {e}")
#             return f"Error generating project-level docs: {e}"

#     # ----------------------------------------
#     # 5️⃣ LEARNING STAGE — Track User Choices
#     # ----------------------------------------
#     def update_user_preferences(self, user_id, accepted, rejected):
#         print(f"\n[LEARNING] 📚 Updating preferences for user: {user_id}")
#         if user_id not in self.user_preferences:
#             self.user_preferences[user_id] = {"accepted": [], "rejected": []}
#         self.user_preferences[user_id]["accepted"].extend(accepted)
#         self.user_preferences[user_id]["rejected"].extend(rejected)
#         print(f"[LEARNING] ✅ Preferences updated — Accepted: {len(accepted)}, Rejected: {len(rejected)}")

#     # ----------------------------------------
#     # 🚀 6️⃣ FULL PIPELINE — Unified AI Developer Flow
#     # ----------------------------------------
#     def run_full_pipeline(self, repo_url, user_id="default_user",
#                           accepted_suggestions=None, rejected_suggestions=None,
#                           doc_format="md", commit_changes=False):
#         print(f"\n🚀 [INIT] Starting full AI Developer pipeline for: {repo_url}")

#         repo_data = self.fetch_repo(repo_url)
#         files = repo_data["files"]
#         if not files:
#             print("⚠️ [INIT] No code files detected — skipping analysis and documentation.")
#             return {
#                 "repo": repo_data["repo_name"],
#                 "analysis": [],
#                 "documentation": [],
#                 "project_doc": "No code files found.",
#                 "user_learning": self.user_preferences.get(user_id, {})
#             }

#         all_analysis, docs_output = [], []

#         for f in files:
#             print(f"\n⚙️ [PROCESS] Working on: {f['file_name']}")
#             analysis = self.analyze_code(f["file_name"], f["content"])
#             analysis["file"] = f["file_name"]
#             all_analysis.append(analysis)

#             if accepted_suggestions:
#                 matched = [s for s in accepted_suggestions if s["file_name"] == f["file_name"]]
#                 if matched:
#                     f["content"] = self.apply_suggestions(f["content"], matched)
#                     if commit_changes:
#                         self.commit_to_github(repo_url, f["file_name"], f["content"], "AI-applied suggestions")

#             doc_text = self.generate_docs(f["file_name"], f["content"], format=doc_format)
#             docs_output.append({"file_name": f["file_name"], "documentation": doc_text})

#         self.update_user_preferences(user_id, accepted_suggestions or [], rejected_suggestions or [])

#         project_doc = self.generate_project_docs(repo_data["repo_name"], files, all_analysis, format=doc_format)

#         print("\n✅ [COMPLETE] AI Developer pipeline executed successfully.")
#         return {
#             "repo": repo_data["repo_name"],
#             "analysis": all_analysis,
#             "documentation": docs_output,
#             "project_doc": project_doc,
#             "user_learning": self.user_preferences.get(user_id, {})
#         }

# # developer_agent.py
# import json, time, requests, os, re
# import google.generativeai as genai
# from autogen import ConversableAgent


# COMMENT_PREFIX = {
#     ".py": "# ",
#     ".js": "// ",
#     ".java": "// ",
#     ".c": "// ",
#     ".cpp": "// ",
#     ".ts": "// ",
#     ".html": "<!-- ",
#     ".css": "/* ",
# }
# COMMENT_SUFFIX = {
#         ".html": " -->",
#         ".css": " */",
# }

# class DeveloperAgent(ConversableAgent):
    

#     def __init__(self, model_name="gemini-2.5-flash"):
#         super().__init__(name="DeveloperAgent", description="Unified AI Developer Agent for CodeZen.")
#         self.model_name = model_name
#         self.model = genai.GenerativeModel(model_name)
#         self.user_preferences = {}
#         print(f"✅ DeveloperAgent initialized with {model_name}")

#     # ----------------------------------------
#     # 1️⃣ PERCEPTION — Fetch repo recursively
#     # ----------------------------------------
#     def fetch_repo(self, repo_url: str):
#         print(f"\n[PERCEPTION] 🌐 Fetching repository structure from: {repo_url}")
#         parts = repo_url.strip("/").split("/")
#         username, reponame = parts[-2], parts[-1].replace(".git", "")
#         meta = requests.get(f"https://api.github.com/repos/{username}/{reponame}").json()
#         default_branch = meta.get("default_branch", "main")
#         print(f"[PERCEPTION] 🪶 Default branch: {default_branch}")

#         tree = requests.get(
#             f"https://api.github.com/repos/{username}/{reponame}/git/trees/{default_branch}?recursive=1"
#         ).json()
#         files, allowed = [], [".py", ".js", ".java", ".ts", ".cpp", ".c", ".html", ".css", ".ipynb"]
#         for item in tree.get("tree", []):
#             if item["type"] == "blob" and any(item["path"].endswith(ext) for ext in allowed):
#                 raw = f"https://raw.githubusercontent.com/{username}/{reponame}/{default_branch}/{item['path']}"
#                 content = requests.get(raw).text
#                 if item["path"].endswith(".ipynb"):
#                     try:
#                         nb = json.loads(content)
#                         cells = ["\n".join(c.get("source", [])) for c in nb.get("cells", []) if c.get("cell_type") == "code"]
#                         content = "\n\n".join(cells)
#                         print(f"📘 [NOTEBOOK] Extracted {len(cells)} code cells from {item['path']}")
#                     except Exception as e:
#                         print(f"⚠️ [NOTEBOOK] Failed to parse notebook {item['path']}: {e}")
#                 files.append({"file_name": item["path"], "content": content})
#         print(f"✅ [PERCEPTION] Repo fetched — {len(files)} code files detected.")
#         return {"repo_name": reponame, "files": files}

#     # ----------------------------------------
#     # 2️⃣ REASONING — Analyze each file
#     # ----------------------------------------
#     def analyze_code(self, file_name, content):
#         print(f"\n[REASONING] 🧩 Analyzing file: {file_name}")
#         prompt = f"""
#         You are an expert AI code reviewer.
#         Analyze the following code and return ONLY valid JSON:
#         {{
#           "issues": ["list of problems or bugs"],
#           "refactors": ["improvement suggestions"],
#           "summary": "short summary of what this file does"
#         }}
#         Code:
#         {content}
#         """
#         try:
#             out = self.model.generate_content(prompt).text.strip()
#             try:
#                 parsed = json.loads(out)
#             except Exception:
#                 match = re.search(r"\{.*\}", out, re.DOTALL)
#                 parsed = json.loads(match.group()) if match else {"issues": [], "refactors": [], "summary": out[:300]}
#             print(f"✅ [REASONING] Parsed for {file_name}")
#             return parsed
#         except Exception as e:
#             print(f"❌ [REASONING] Error: {e}")
#             return {"file": file_name, "issues": [], "refactors": [], "summary": str(e)}

#     # ----------------------------------------
#     # 3️⃣ ACTION — Apply suggestions (simulated)
#     # -----------------------------------------
#     def _parse_repo(self, repo_url):
#         parts = repo_url.strip("/").split("/")
#         return parts[-2], parts[-1].replace(".git", "")

#     def _github_headers(self):
#         token = os.getenv("GITHUB_TOKEN")
#         if not token:
#             raise Exception("Missing GITHUB_TOKEN environment variable.")
#         return {
#             "Authorization": f"token {token}",
#             "Accept": "application/vnd.github+json",
#         }

#     def _get_default_branch(self, username, reponame):
#         resp = requests.get(f"https://api.github.com/repos/{username}/{reponame}", headers=self._github_headers())
#         resp.raise_for_status()
#         return resp.json().get("default_branch", "main")

#     def _ensure_feature_branch(self, username, reponame, base_branch, feature_branch):
#         ref_url = f"https://api.github.com/repos/{username}/{reponame}/git/ref/heads/{base_branch}"
#         ref_data = requests.get(ref_url, headers=self._github_headers()).json()
#         base_sha = ref_data["object"]["sha"]

#         create_resp = requests.post(
#             f"https://api.github.com/repos/{username}/{reponame}/git/refs",
#             headers=self._github_headers(),
#             json={"ref": f"refs/heads/{feature_branch}", "sha": base_sha},
#         )
#         if create_resp.status_code == 201:
#             print(f"✅ [ACTION] Created new branch: {feature_branch}")
#         elif create_resp.status_code == 422:
#             print(f"⚠️ [ACTION] Branch {feature_branch} already exists (reusing).")
#         else:
#             print(f"❌ [ACTION] Failed to create branch: {create_resp.text}")

#     def _get_file_content_and_sha(self, username, reponame, branch, path):
#         resp = requests.get(
#             f"https://api.github.com/repos/{username}/{reponame}/contents/{path}?ref={branch}",
#             headers=self._github_headers()
#         )
#         resp.raise_for_status()
#         data = resp.json()
#         import base64
#         content = base64.b64decode(data["content"]).decode("utf-8", errors="ignore")
#         return content, data["sha"]

#     def _commit_file_update(self, username, reponame, branch, path, new_content, commit_message, sha):
#         import base64
#         encoded = base64.b64encode(new_content.encode("utf-8")).decode("utf-8")
#         resp = requests.put(
#             f"https://api.github.com/repos/{username}/{reponame}/contents/{path}",
#             headers=self._github_headers(),
#             json={
#                 "message": commit_message,
#                 "content": encoded,
#                 "branch": branch,
#                 "sha": sha,
#             },
#         )
#         resp.raise_for_status()
#         print(f"✅ [ACTION] Committed {path} → {branch}")

#     # ----------------------------------------
#     def apply_suggestions(self, file_path, original_content, suggestions):
#         from datetime import datetime
#         ext = "." + file_path.split(".")[-1]
#         prefix = COMMENT_PREFIX.get(ext, "# ")
#         suffix = COMMENT_SUFFIX.get(ext, "")
#         timestamp = datetime.utcnow().isoformat()

#         summary_block = "\n".join(
#             [f"{prefix}Applied suggestion: {s.get('comment','').strip()}{suffix}" for s in suggestions]
#         )
#         header = f"{prefix}CodeZen Fixes ({timestamp}){suffix}\n{summary_block}\n\n"
#         return header + original_content


#     # ----------------------------------------
#     # 4️⃣ SMART PROJECT-LEVEL DOCS (Mermaid diagram)
#     # ----------------------------------------
#     def generate_project_docs(self, repo_name, files, analyses, format="md"):
#         print(f"\n[PROJECT DOCS] 🏗️ Generating {format.upper()} documentation with smart diagram for: {repo_name}")

#         # Simple pattern-based role detection
#         modules = {"frontend": [], "backend": [], "database": [], "ml": [], "api": []}
#         for f in files:
#             path = f["file_name"].lower()
#             if any(k in path for k in ["html", "css", "js", "react"]):
#                 modules["frontend"].append(path)
#             elif any(k in path for k in ["api", "server", "flask", "fastapi", "node", "spring"]):
#                 modules["backend"].append(path)
#             elif any(k in path for k in ["sql", "db", "database", "model"]):
#                 modules["database"].append(path)
#             elif any(k in path for k in ["ml", "train", "model", "predict"]):
#                 modules["ml"].append(path)
#             else:
#                 modules["api"].append(path)

#         # Dynamically build a mermaid diagram structure
#         mermaid = "graph TD\n"
#         mermaid += "    A[User] --> B[Frontend]\n" if modules["frontend"] else ""
#         if modules["backend"]:
#             mermaid += "    B --> C[Backend/API]\n" if modules["frontend"] else "    A --> C[Backend]\n"
#         if modules["ml"]:
#             mermaid += "    C --> D[ML Model]\n"
#         if modules["database"]:
#             mermaid += "    C --> E[Database]\n"
#         mermaid += "    D --> C\n    E --> C\n"

#         summarized = "\n".join(
#             [f"- {a.get('file', 'N/A')}: {a.get('summary', '')}" for a in analyses if a.get("summary")]
#         )

#         combined_code = "\n".join([f"### FILE: {f['file_name']}\n{f['content']}" for f in files])

#         project_prompt = f"""
#         You are a senior software architect.
#         Create a **Technical Design Document (TTD)** for the project **{repo_name}**.

#         The doc must include:
#         1. **Overview & Purpose**
#         2. **Architecture Diagram (Mermaid)**
#         3. **Modules / Components Overview**
#         4. **Data & Control Flow**
#         5. **Technologies Used**
#         6. **Execution Flow**
#         7. **Limitations / Future Work**
#         8. **Conclusion**

#         Include this Mermaid diagram exactly:

#         ```mermaid
#         {mermaid}
#         ```

#         ### Summarized Analysis
#         {summarized}

#         ### Code Snapshot
#         {combined_code}
#         """

#         try:
#             res = self.model.generate_content(project_prompt)
#             print(f"[PROJECT DOCS] ✅ Project-level doc generated with smart diagram.")
#             return res.text.strip()
#         except Exception as e:
#             print(f"❌ [PROJECT DOCS] Error: {e}")
#             return f"Error generating project docs: {e}"

#     # ----------------------------------------
#     # 5️⃣ LEARNING — Track user feedback
#     # ----------------------------------------
#     def update_user_preferences(self, user_id, accepted, rejected):
#         if user_id not in self.user_preferences:
#             self.user_preferences[user_id] = {"accepted": [], "rejected": []}
#         self.user_preferences[user_id]["accepted"].extend(accepted)
#         self.user_preferences[user_id]["rejected"].extend(rejected)

#     # ----------------------------------------
#     # 6️⃣ FULL PIPELINE — Per-file analysis + unified doc
#     # ----------------------------------------
# # 6️⃣ FULL PIPELINE — Per-file analysis + unified doc + commit
# # ----------------------------------------
# def run_full_pipeline(
#     self,
#     repo_url,
#     user_id="default_user",
#     accepted_suggestions=None,
#     rejected_suggestions=None,
#     doc_format="md",
#     commit_changes=False
# ):
#     print(f"\n🚀 [INIT] Running AI Developer pipeline for: {repo_url}")

#     # -------------------------------
#     # 🧠 PERCEPTION — Fetch Repo Files
#     # -------------------------------
#     repo = self.fetch_repo(repo_url)
#     files = repo["files"]
#     if not files:
#         return {
#             "repo": repo["repo_name"],
#             "analysis": [],
#             "project_doc": "No code files found.",
#         }

#     # -------------------------------
#     # 🧩 REASONING — Analyze Each File
#     # -------------------------------
#     analyses = []
#     for f in files:
#         a = self.analyze_code(f["file_name"], f["content"])
#         a["file"] = f["file_name"]
#         analyses.append(a)

#     # -------------------------------
#     # ⚙️ ACTION — Apply Suggestions + Commit (if accepted)
#     # -------------------------------
#     if commit_changes and accepted_suggestions:
#         print("\n[ACTION] 🧩 Applying developer-approved suggestions and committing changes...")
#         username, reponame = self._parse_repo(repo_url)
#         base_branch = self._get_default_branch(username, reponame)
#         feature_branch = f"codezen/fixes-{int(time.time())}"
#         self._ensure_feature_branch(username, reponame, base_branch, feature_branch)

#         # Group accepted suggestions per file
#         grouped = {}
#         for s in accepted_suggestions:
#             if s.get("accepted"):
#                 grouped.setdefault(s["file_name"], []).append(s)

#         for file_name, suggs in grouped.items():
#             try:
#                 current_content, sha = self._get_file_content_and_sha(
#                     username, reponame, feature_branch, file_name
#                 )
#                 updated_content = self.apply_suggestions(file_name, current_content, suggs)
#                 message = f"AI-Fix: Applied {len(suggs)} developer-approved suggestion(s) to {file_name}"
#                 self._commit_file_update(
#                     username, reponame, feature_branch, file_name, updated_content, message, sha
#                 )
#             except Exception as e:
#                 print(f"❌ [ACTION] Failed to commit {file_name}: {e}")
#         print(f"✅ [ACTION] All accepted suggestions committed to branch: {feature_branch}")
#     else:
#         print("[ACTION] ⚪ Skipping commit — No accepted suggestions or commit disabled.")

#     # -------------------------------
#     # 🧾 LEARNING — Track Feedback
#     # -------------------------------
#     self.update_user_preferences(user_id, accepted_suggestions or [], rejected_suggestions or [])

#     # -------------------------------
#     # 📘 DOCUMENTATION — Generate Project-Level Docs
#     # -------------------------------
#     project_doc = self.generate_project_docs(
#         repo["repo_name"], files, analyses, format=doc_format
#     )

#     print("✅ [COMPLETE] AI Developer pipeline executed successfully.")
#     return {
#         "repo": repo["repo_name"],
#         "analysis": analyses,
#         "project_doc": project_doc,
#         "user_learning": self.user_preferences.get(user_id, {}),
#     }


# import json, time, requests, os, re
# import google.generativeai as genai
# from autogen import ConversableAgent

# # ----------------------------------------
# # Comment Syntax Configurations
# # ----------------------------------------
# COMMENT_PREFIX = {
#     ".py": "# ",
#     ".js": "// ",
#     ".java": "// ",
#     ".c": "// ",
#     ".cpp": "// ",
#     ".ts": "// ",
#     ".html": "<!-- ",
#     ".css": "/* ",
# }
# COMMENT_SUFFIX = {
#     ".html": " -->",
#     ".css": " */",
# }


# class DeveloperAgent(ConversableAgent):
#     def __init__(self, model_name="gemini-2.5-flash"):
#         super().__init__(name="DeveloperAgent", description="Unified AI Developer Agent for CodeZen.")
#         self.model_name = model_name
#         self.model = genai.GenerativeModel(model_name)
#         self.user_preferences = {}
#         print(f"✅ DeveloperAgent initialized with {model_name}")

#     # ----------------------------------------
#     # 1️⃣ PERCEPTION — Fetch repo recursively
#     # ----------------------------------------
#     def fetch_repo(self, repo_url: str):
#         print(f"\n[PERCEPTION] 🌐 Fetching repository structure from: {repo_url}")
#         parts = repo_url.strip("/").split("/")
#         username, reponame = parts[-2], parts[-1].replace(".git", "")
#         meta = requests.get(f"https://api.github.com/repos/{username}/{reponame}").json()
#         default_branch = meta.get("default_branch", "main")
#         print(f"[PERCEPTION] 🪶 Default branch: {default_branch}")

#         tree = requests.get(
#             f"https://api.github.com/repos/{username}/{reponame}/git/trees/{default_branch}?recursive=1"
#         ).json()
#         files, allowed = [], [".py", ".js", ".java", ".ts", ".cpp", ".c", ".html", ".css", ".ipynb"]
#         for item in tree.get("tree", []):
#             if item["type"] == "blob" and any(item["path"].endswith(ext) for ext in allowed):
#                 raw = f"https://raw.githubusercontent.com/{username}/{reponame}/{default_branch}/{item['path']}"
#                 content = requests.get(raw).text
#                 if item["path"].endswith(".ipynb"):
#                     try:
#                         nb = json.loads(content)
#                         cells = [
#                             "\n".join(c.get("source", []))
#                             for c in nb.get("cells", [])
#                             if c.get("cell_type") == "code"
#                         ]
#                         content = "\n\n".join(cells)
#                         print(f"📘 [NOTEBOOK] Extracted {len(cells)} code cells from {item['path']}")
#                     except Exception as e:
#                         print(f"⚠️ [NOTEBOOK] Failed to parse notebook {item['path']}: {e}")
#                 files.append({"file_name": item["path"], "content": content})
#         print(f"✅ [PERCEPTION] Repo fetched — {len(files)} code files detected.")
#         return {"repo_name": reponame, "files": files}

#     # ----------------------------------------
#     # 2️⃣ REASONING — Analyze each file
#     # ----------------------------------------
#     def analyze_code(self, file_name, content):
#         print(f"\n[REASONING] 🧩 Analyzing file: {file_name}")
#         prompt = f"""
#         You are an expert AI code reviewer.
#         Analyze the following code and return ONLY valid JSON:
#         {{
#           "issues": ["list of problems or bugs"],
#           "refactors": ["improvement suggestions"],
#           "summary": "short summary of what this file does"
#         }}
#         Code:
#         {content}
#         """
#         try:
#             out = self.model.generate_content(prompt).text.strip()
#             try:
#                 parsed = json.loads(out)
#             except Exception:
#                 match = re.search(r"\{.*\}", out, re.DOTALL)
#                 parsed = (
#                     json.loads(match.group())
#                     if match
#                     else {"issues": [], "refactors": [], "summary": out[:300]}
#                 )
#             print(f"✅ [REASONING] Parsed for {file_name}")
#             return parsed
#         except Exception as e:
#             print(f"❌ [REASONING] Error: {e}")
#             return {"file": file_name, "issues": [], "refactors": [], "summary": str(e)}

#     # ----------------------------------------
#     # 3️⃣ ACTION — GitHub Commit Helpers
#     # ----------------------------------------
#     def _parse_repo(self, repo_url):
#         parts = repo_url.strip("/").split("/")
#         return parts[-2], parts[-1].replace(".git", "")

#     def _github_headers(self):
#         token = os.getenv("GITHUB_TOKEN")
#         if not token:
#             raise Exception("Missing GITHUB_TOKEN environment variable.")
#         return {"Authorization": f"token {token}", "Accept": "application/vnd.github+json"}

#     def _get_default_branch(self, username, reponame):
#         resp = requests.get(f"https://api.github.com/repos/{username}/{reponame}", headers=self._github_headers())
#         resp.raise_for_status()
#         return resp.json().get("default_branch", "main")

#     def _ensure_feature_branch(self, username, reponame, base_branch, feature_branch):
#         ref_url = f"https://api.github.com/repos/{username}/{reponame}/git/ref/heads/{base_branch}"
#         ref_data = requests.get(ref_url, headers=self._github_headers()).json()
#         base_sha = ref_data["object"]["sha"]

#         create_resp = requests.post(
#             f"https://api.github.com/repos/{username}/{reponame}/git/refs",
#             headers=self._github_headers(),
#             json={"ref": f"refs/heads/{feature_branch}", "sha": base_sha},
#         )
#         if create_resp.status_code == 201:
#             print(f"✅ [ACTION] Created new branch: {feature_branch}")
#         elif create_resp.status_code == 422:
#             print(f"⚠️ [ACTION] Branch {feature_branch} already exists (reusing).")
#         else:
#             print(f"❌ [ACTION] Failed to create branch: {create_resp.text}")

#     def _get_file_content_and_sha(self, username, reponame, branch, path):
#         resp = requests.get(
#             f"https://api.github.com/repos/{username}/{reponame}/contents/{path}?ref={branch}",
#             headers=self._github_headers(),
#         )
#         resp.raise_for_status()
#         data = resp.json()
#         import base64
#         content = base64.b64decode(data["content"]).decode("utf-8", errors="ignore")
#         return content, data["sha"]

#     def _commit_file_update(self, username, reponame, branch, path, new_content, commit_message, sha):
#         import base64
#         encoded = base64.b64encode(new_content.encode("utf-8")).decode("utf-8")
#         resp = requests.put(
#             f"https://api.github.com/repos/{username}/{reponame}/contents/{path}",
#             headers=self._github_headers(),
#             json={"message": commit_message, "content": encoded, "branch": branch, "sha": sha},
#         )
#         resp.raise_for_status()
#         print(f"✅ [ACTION] Committed {path} → {branch}")

#     # ----------------------------------------
#     def apply_suggestions(self, file_path, original_content, suggestions):
#         from datetime import datetime
#         ext = "." + file_path.split(".")[-1]
#         prefix = COMMENT_PREFIX.get(ext, "# ")
#         suffix = COMMENT_SUFFIX.get(ext, "")
#         timestamp = datetime.utcnow().isoformat()

#         summary_block = "\n".join(
#             [f"{prefix}Applied suggestion: {s.get('comment','').strip()}{suffix}" for s in suggestions]
#         )
#         header = f"{prefix}CodeZen Fixes ({timestamp}){suffix}\n{summary_block}\n\n"
#         return header + original_content

#     # ----------------------------------------
#     # 4️⃣ SMART PROJECT-LEVEL DOCS (Mermaid diagram)
#     # ----------------------------------------
#     def generate_project_docs(self, repo_name, files, analyses, format="md"):
#         print(f"\n[PROJECT DOCS] 🏗️ Generating {format.upper()} documentation with smart diagram for: {repo_name}")

#         modules = {"frontend": [], "backend": [], "database": [], "ml": [], "api": []}
#         for f in files:
#             path = f["file_name"].lower()
#             if any(k in path for k in ["html", "css", "js", "react"]):
#                 modules["frontend"].append(path)
#             elif any(k in path for k in ["api", "server", "flask", "fastapi", "node", "spring"]):
#                 modules["backend"].append(path)
#             elif any(k in path for k in ["sql", "db", "database", "model"]):
#                 modules["database"].append(path)
#             elif any(k in path for k in ["ml", "train", "model", "predict"]):
#                 modules["ml"].append(path)
#             else:
#                 modules["api"].append(path)

#         mermaid = "graph TD\n"
#         mermaid += "    A[User] --> B[Frontend]\n" if modules["frontend"] else ""
#         if modules["backend"]:
#             mermaid += "    B --> C[Backend/API]\n" if modules["frontend"] else "    A --> C[Backend]\n"
#         if modules["ml"]:
#             mermaid += "    C --> D[ML Model]\n"
#         if modules["database"]:
#             mermaid += "    C --> E[Database]\n"
#         mermaid += "    D --> C\n    E --> C\n"

#         summarized = "\n".join([f"- {a.get('file', 'N/A')}: {a.get('summary', '')}" for a in analyses if a.get("summary")])
#         combined_code = "\n".join([f"### FILE: {f['file_name']}\n{f['content']}" for f in files])

#         project_prompt = f"""
#         You are a senior software architect.
#         Create a **Technical Design Document (TTD)** for the project **{repo_name}**.

#         Include this Mermaid diagram exactly:

#         ```mermaid
#         {mermaid}
#         ```

#         ### Summarized Analysis
#         {summarized}

#         ### Code Snapshot
#         {combined_code}
#         """

#         try:
#             res = self.model.generate_content(project_prompt)
#             print(f"[PROJECT DOCS] ✅ Project-level doc generated with smart diagram.")
#             return res.text.strip()
#         except Exception as e:
#             print(f"❌ [PROJECT DOCS] Error: {e}")
#             return f"Error generating project docs: {e}"

#     # ----------------------------------------
#     # 5️⃣ LEARNING — Track user feedback
#     # ----------------------------------------
#     # ----------------------------------------

#     def update_user_preferences(self, user_id, accepted, rejected):
#         memory_file = ".codezen_memory.json"

#         # Load existing memory
#         if os.path.exists(memory_file):
#             try:
#                 with open(memory_file, "r", encoding="utf-8") as f:
#                     memory_data = json.load(f)
#             except Exception:
#                 memory_data = {}
#         else:
#             memory_data = {}

#         # Ensure user entry
#         if user_id not in memory_data:
#             memory_data[user_id] = {"accepted": [], "rejected": []}

#         # Append new learning data
#         accepted_data = [
#             {"file_name": s.get("file_name"), "comment": s.get("comment"), "timestamp": time.time()}
#             for s in accepted or []
#         ]
#         rejected_data = [
#             {"file_name": s.get("file_name"), "comment": s.get("comment"), "timestamp": time.time()}
#             for s in rejected or []
#         ]

#         memory_data[user_id]["accepted"].extend(accepted_data)
#         memory_data[user_id]["rejected"].extend(rejected_data)

#         # Save updated memory
#         with open(memory_file, "w", encoding="utf-8") as f:
#             json.dump(memory_data, f, indent=2)

#         self.user_preferences[user_id] = memory_data[user_id]
#         print(f"🧠 [LEARNING] Updated user preferences for '{user_id}' — {len(accepted_data)} accepted, {len(rejected_data)} rejected suggestions.")


#     # ----------------------------------------
#     # 6️⃣ FULL PIPELINE — Per-file analysis + unified doc + commit
#     # ----------------------------------------
#     def run_full_pipeline(
#         self,
#         repo_url,
#         user_id="default_user",
#         accepted_suggestions=None,
#         rejected_suggestions=None,
#         doc_format="md",
#         commit_changes=False,
#     ):
#         print(f"\n🚀 [INIT] Running AI Developer pipeline for: {repo_url}")

#         # 🧠 PERCEPTION
#         repo = self.fetch_repo(repo_url)
#         files = repo["files"]
#         if not files:
#             return {"repo": repo["repo_name"], "analysis": [], "project_doc": "No code files found."}

#         # 🧩 REASONING
#         analyses = []
#         for f in files:
#             a = self.analyze_code(f["file_name"], f["content"])
#             a["file"] = f["file_name"]
#             analyses.append(a)

#         # ⚙️ ACTION — Commit if accepted
#         if commit_changes and accepted_suggestions:
#             print("\n[ACTION] 🧩 Applying developer-approved suggestions and committing changes...")
#             username, reponame = self._parse_repo(repo_url)
#             base_branch = self._get_default_branch(username, reponame)
#             feature_branch = f"codezen/fixes-{int(time.time())}"
#             self._ensure_feature_branch(username, reponame, base_branch, feature_branch)

#             grouped = {}
#             for s in accepted_suggestions:
#                 if s.get("accepted"):
#                     grouped.setdefault(s["file_name"], []).append(s)

#             for file_name, suggs in grouped.items():
#                 try:
#                     current_content, sha = self._get_file_content_and_sha(username, reponame, feature_branch, file_name)
#                     updated_content = self.apply_suggestions(file_name, current_content, suggs)
#                     message = f"AI-Fix: Applied {len(suggs)} developer-approved suggestion(s) to {file_name}"
#                     self._commit_file_update(username, reponame, feature_branch, file_name, updated_content, message, sha)
#                 except Exception as e:
#                     print(f"❌ [ACTION] Failed to commit {file_name}: {e}")
#             print(f"✅ [ACTION] All accepted suggestions committed to branch: {feature_branch}")
#         else:
#             print("[ACTION] ⚪ Skipping commit — No accepted suggestions or commit disabled.")

#         # 🧾 LEARNING
#         self.update_user_preferences(user_id, accepted_suggestions or [], rejected_suggestions or [])

#         # 📘 DOCUMENTATION
#         project_doc = self.generate_project_docs(repo["repo_name"], files, analyses, format=doc_format)

#         print("✅ [COMPLETE] AI Developer pipeline executed successfully.")
#         return {
#             "repo": repo["repo_name"],
#             "analysis": analyses,
#             "project_doc": project_doc,
#             "user_learning": self.user_preferences.get(user_id, {}),
#         }


# import json, time, requests, os, re
# import google.generativeai as genai
# from autogen import ConversableAgent

# # ----------------------------------------
# # Comment Syntax Configurations
# # ----------------------------------------
# COMMENT_PREFIX = {
#     ".py": "# ",
#     ".js": "// ",
#     ".java": "// ",
#     ".c": "// ",
#     ".cpp": "// ",
#     ".ts": "// ",
#     ".html": "<!-- ",
#     ".css": "/* ",
# }
# COMMENT_SUFFIX = {
#     ".html": " -->",
#     ".css": " */",
# }

# # developer_agent.py (only the NEW/CHANGED parts shown)

# class DeveloperAgent(ConversableAgent):
#     # ... keep your __init__, helpers, analyze_code, commit helpers, generate_project_docs, update_user_preferences

#     # --- NEW: Analysis-only (Perception + Reasoning) ---
#     def run_analysis_only(self, repo_url):
#         print(f"\n🚀 [INIT] Running ANALYSIS-ONLY for: {repo_url}")
#         repo = self.fetch_repo(repo_url)
#         files = repo["files"]
#         if not files:
#             return {
#                 "repo": repo["repo_name"],
#                 "analysis": [],
#                 "project_doc": None
#             }

#         analyses = []
#         for f in files:
#             a = self.analyze_code(f["file_name"], f["content"])
#             a["file"] = f["file_name"]
#             analyses.append(a)

#         # Return ONLY the analysis to the frontend
#         return {
#             "repo": repo["repo_name"],
#             "analysis": analyses,
#             "project_doc": None
#         }

#     # --- NEW: Apply Changes (optional) + Generate Docs ---
#     def apply_changes_and_generate_docs(
#         self,
#         repo_url,
#         user_id="default_user",
#         accepted_suggestions=None,
#         rejected_suggestions=None,
#         doc_format="md",
#         commit_changes=False,
#     ):
#         print(f"\n🚀 [INIT] APPLY-CHANGES for: {repo_url}")
#         # Re-fetch repo and files (source of truth)
#         repo = self.fetch_repo(repo_url)
#         files = repo["files"]

#         # Optional commit phase
#         if commit_changes and accepted_suggestions:
#             print("\n[ACTION] 🧩 Applying developer-approved suggestions and committing changes...")
#             username, reponame = self._parse_repo(repo_url)
#             base_branch = self._get_default_branch(username, reponame)
#             feature_branch = f"codezen/fixes-{int(time.time())}"
#             self._ensure_feature_branch(username, reponame, base_branch, feature_branch)

#             grouped = {}
#             for s in accepted_suggestions:
#                 if (isinstance(s, dict) and s.get("accepted")) or getattr(s, "accepted", False):
#                     fname = s.get("file_name") if isinstance(s, dict) else s.file_name
#                     grouped.setdefault(fname, []).append(
#                         s if isinstance(s, dict) else s.model_dump()
#                     )

#             for file_name, suggs in grouped.items():
#                 try:
#                     current_content, sha = self._get_file_content_and_sha(username, reponame, feature_branch, file_name)
#                     updated_content = self.apply_suggestions(file_name, current_content, suggs)
#                     message = f"AI-Fix: Applied {len(suggs)} developer-approved suggestion(s) to {file_name}"
#                     self._commit_file_update(username, reponame, feature_branch, file_name, updated_content, message, sha)
#                 except Exception as e:
#                     print(f"❌ [ACTION] Failed to commit {file_name}: {e}")
#             print(f"✅ [ACTION] All accepted suggestions committed to branch: {feature_branch}")
#         else:
#             print("[ACTION] ⚪ Skipping commit — No accepted suggestions or commit disabled.")

#         # Learning loop
#         self.update_user_preferences(user_id, accepted_suggestions or [], rejected_suggestions or [])

#         # Fresh analyses (optional) before docs; for now reuse previous files
#         analyses = []
#         for f in files:
#             a = self.analyze_code(f["file_name"], f["content"])
#             a["file"] = f["file_name"]
#             analyses.append(a)

#         # Generate docs
#         project_doc = self.generate_project_docs(repo["repo_name"], files, analyses, format=doc_format)

#         return {
#             "repo": repo["repo_name"],
#             "analysis": analyses,
#             "project_doc": project_doc,
#             "user_learning": self.user_preferences.get(user_id, {}),
#         }



# import json, time, requests, os, re
# import google.generativeai as genai
# from autogen import ConversableAgent

# # ----------------------------------------
# # Comment Syntax Configurations
# # ----------------------------------------
# COMMENT_PREFIX = {
#     ".py": "# ",
#     ".js": "// ",
#     ".java": "// ",
#     ".c": "// ",
#     ".cpp": "// ",
#     ".ts": "// ",
#     ".html": "<!-- ",
#     ".css": "/* ",
# }
# COMMENT_SUFFIX = {
#     ".html": " -->",
#     ".css": " */",
# }

# # developer_agent.py (only the NEW/CHANGED parts shown)

# # 
# class DeveloperAgent(ConversableAgent):
#     def __init__(self, name="DeveloperAgent", description="CodeZen core developer agent", model="gemini-1.5-pro"):
#         super().__init__(name=name, description=description, model=model)
#         self.user_preferences = {}  # for adaptive learning
#         self.github_token = None

#     # --- NEW: Analysis-only (Perception + Reasoning) ---
#     def run_analysis_only(self, repo_url):
#         print(f"\n🚀 [INIT] Running ANALYSIS-ONLY for: {repo_url}")
#         repo = self.fetch_repo(repo_url)
#         files = repo["files"]
#         if not files:
#             return {
#                 "repo": repo["repo_name"],
#                 "analysis": [],
#                 "project_doc": None
#             }

#         analyses = []
#         for f in files:
#             a = self.analyze_code(f["file_name"], f["content"])
#             a["file"] = f["file_name"]
#             analyses.append(a)

#         # Return ONLY the analysis to the frontend
#         return {
#             "repo": repo["repo_name"],
#             "analysis": analyses,
#             "project_doc": None
#         }

#     # --- NEW: Apply Changes (optional) + Generate Docs ---
#     def apply_changes_and_generate_docs(
#         self,
#         repo_url,
#         user_id="default_user",
#         accepted_suggestions=None,
#         rejected_suggestions=None,
#         doc_format="md",
#         commit_changes=False,
#     ):
#         print(f"\n🚀 [INIT] APPLY-CHANGES for: {repo_url}")
#         # Re-fetch repo and files (source of truth)
#         repo = self.fetch_repo(repo_url)
#         files = repo["files"]

#         # Optional commit phase
#         if commit_changes and accepted_suggestions:
#             print("\n[ACTION] 🧩 Applying developer-approved suggestions and committing changes...")
#             username, reponame = self._parse_repo(repo_url)
#             base_branch = self._get_default_branch(username, reponame)
#             feature_branch = f"codezen/fixes-{int(time.time())}"
#             self._ensure_feature_branch(username, reponame, base_branch, feature_branch)

#             grouped = {}
#             for s in accepted_suggestions:
#                 if (isinstance(s, dict) and s.get("accepted")) or getattr(s, "accepted", False):
#                     fname = s.get("file_name") if isinstance(s, dict) else s.file_name
#                     grouped.setdefault(fname, []).append(
#                         s if isinstance(s, dict) else s.model_dump()
#                     )

#             for file_name, suggs in grouped.items():
#                 try:
#                     current_content, sha = self._get_file_content_and_sha(username, reponame, feature_branch, file_name)
#                     updated_content = self.apply_suggestions(file_name, current_content, suggs)
#                     message = f"AI-Fix: Applied {len(suggs)} developer-approved suggestion(s) to {file_name}"
#                     self._commit_file_update(username, reponame, feature_branch, file_name, updated_content, message, sha)
#                 except Exception as e:
#                     print(f"❌ [ACTION] Failed to commit {file_name}: {e}")
#             print(f"✅ [ACTION] All accepted suggestions committed to branch: {feature_branch}")
#         else:
#             print("[ACTION] ⚪ Skipping commit — No accepted suggestions or commit disabled.")

#         # Learning loop
#         self.update_user_preferences(user_id, accepted_suggestions or [], rejected_suggestions or [])

#         # Fresh analyses (optional) before docs; for now reuse previous files
#         analyses = []
#         for f in files:
#             a = self.analyze_code(f["file_name"], f["content"])
#             a["file"] = f["file_name"]
#             analyses.append(a)

#         # Generate docs
#         project_doc = self.generate_project_docs(repo["repo_name"], files, analyses, format=doc_format)

#         return {
#             "repo": repo["repo_name"],
#             "analysis": analyses,
#             "project_doc": project_doc,
#             "user_learning": self.user_preferences.get(user_id, {}),
#         }

# import json, time, requests, os, re, base64
# import google.generativeai as genai
# from autogen import ConversableAgent

# # ----------------------------------------
# # Comment Syntax Configurations
# # ----------------------------------------
# COMMENT_PREFIX = {
#     ".py": "# ", ".js": "// ", ".java": "// ", ".c": "// ", ".cpp": "// ",
#     ".ts": "// ", ".html": "<!-- ", ".css": "/* ",
# }
# COMMENT_SUFFIX = {".html": " -->", ".css": " */"}


# class DeveloperAgent(ConversableAgent):
#     # 🔵 STEP 1: UPDATED INIT AND TOKEN HANDLING METHODS
#     def __init__(self, model_name="gemini-2.5-flash"):
#         super().__init__(name="DeveloperAgent", description="Unified AI Developer Agent for CodeZen.")
#         self.model_name = model_name
#         self.model = genai.GenerativeModel(model_name)
#         self.user_preferences = {}
#         self.github_token = None  # ✅ stores user token dynamically
#         print(f"✅ DeveloperAgent initialized with {model_name}")

#     def set_dynamic_token(self, token: str):
#         """Store GitHub token from frontend OAuth for current user session."""
#         self.github_token = token
#         print("🔐 Dynamic GitHub token set successfully.")

#     # ----------------------------------------
#     # GitHub header helper (with token check)
#     # ----------------------------------------
#     def _github_headers(self):
#         """Use dynamic token from OAuth if available, otherwise fall back to env variable."""
#         token = getattr(self, "github_token", None) or os.getenv("GITHUB_TOKEN")
#         if not token:
#             raise Exception("Missing GitHub token (neither dynamic nor environment).")
#         return {"Authorization": f"token {token}", "Accept": "application/vnd.github+json"}

#     # ----------------------------------------
#     # Token validation helper (NEW)
#     # ----------------------------------------
#     def _validate_github_token(self):
#         """Validate current GitHub token; returns True if valid, False if expired/invalid."""
#         try:
#             headers = self._github_headers()
#             response = requests.get("https://api.github.com/user", headers=headers)
#             if response.status_code == 200:
#                 return True
#             elif response.status_code in [401, 403]:
#                 print("⚠️ [AUTH] GitHub token is invalid or expired.")
#                 self.github_token = None
#                 return False
#             else:
#                 print(f"⚠️ [AUTH] Unexpected GitHub API status: {response.status_code}")
#                 return False
#         except Exception as e:
#             print(f"❌ [AUTH] Error validating token: {e}")
#             return False

#     # ----------------------------------------
#     # 1️⃣ PERCEPTION — Fetch repo recursively
#     # ----------------------------------------
#     def fetch_repo(self, repo_url: str):
#         print(f"\n[PERCEPTION] 🌐 Fetching repository structure from: {repo_url}")
#         parts = repo_url.strip("/").split("/")
#         username, reponame = parts[-2], parts[-1].replace(".git", "")
#         meta = requests.get(f"https://api.github.com/repos/{username}/{reponame}").json()
#         default_branch = meta.get("default_branch", "main")
#         print(f"[PERCEPTION] 🪶 Default branch: {default_branch}")

#         tree = requests.get(
#             f"https://api.github.com/repos/{username}/{reponame}/git/trees/{default_branch}?recursive=1"
#         ).json()
#         files, allowed = [], [".py", ".js", ".java", ".ts", ".cpp", ".c", ".html", ".css", ".ipynb"]
#         for item in tree.get("tree", []):
#             if item["type"] == "blob" and any(item["path"].endswith(ext) for ext in allowed):
#                 raw = f"https://raw.githubusercontent.com/{username}/{reponame}/{default_branch}/{item['path']}"
#                 content = requests.get(raw).text
#                 if item["path"].endswith(".ipynb"):
#                     try:
#                         nb = json.loads(content)
#                         cells = [
#                             "\n".join(c.get("source", []))
#                             for c in nb.get("cells", [])
#                             if c.get("cell_type") == "code"
#                         ]
#                         content = "\n\n".join(cells)
#                         print(f"📘 [NOTEBOOK] Extracted {len(cells)} code cells from {item['path']}")
#                     except Exception as e:
#                         print(f"⚠️ [NOTEBOOK] Failed to parse notebook {item['path']}: {e}")
#                 files.append({"file_name": item["path"], "content": content})
#         print(f"✅ [PERCEPTION] Repo fetched — {len(files)} code files detected.")
#         return {"repo_name": reponame, "files": files}

#     # ----------------------------------------
#     # 2️⃣ REASONING — Analyze each file with Gemini
#     # ----------------------------------------
#     def analyze_code(self, file_name, content):
#         print(f"\n[REASONING] 🧩 Analyzing file: {file_name}")
#         prompt = f"""
#         You are an expert AI code reviewer.
#         Analyze the following code and return ONLY valid JSON:
#         {{
#           "issues": ["list of problems or bugs"],
#           "refactors": ["improvement suggestions"],
#           "summary": "short summary of what this file does"
#         }}
#         Code:
#         {content}
#         """
#         try:
#             out = self.model.generate_content(prompt).text.strip()
#             try:
#                 parsed = json.loads(out)
#             except Exception:
#                 match = re.search(r"\{.*\}", out, re.DOTALL)
#                 parsed = (
#                     json.loads(match.group())
#                     if match
#                     else {"issues": [], "refactors": [], "summary": out[:300]}
#                 )
#             print(f"✅ [REASONING] Parsed for {file_name}")
#             return parsed
#         except Exception as e:
#             print(f"❌ [REASONING] Error: {e}")
#             return {"file": file_name, "issues": [], "refactors": [], "summary": str(e)}

#     # ----------------------------------------
#     # 3️⃣ ACTION — GitHub Commit Helpers
#     # ----------------------------------------
#     def _parse_repo(self, repo_url):
#         parts = repo_url.strip("/").split("/")
#         return parts[-2], parts[-1].replace(".git", "")

#     def _get_default_branch(self, username, reponame):
#         if not self._validate_github_token():
#             raise Exception("GitHub token expired or invalid. Please re-login via frontend.")
#         resp = requests.get(f"https://api.github.com/repos/{username}/{reponame}", headers=self._github_headers())
#         resp.raise_for_status()
#         return resp.json().get("default_branch", "main")

#     def _ensure_feature_branch(self, username, reponame, base_branch, feature_branch):
#         if not self._validate_github_token():
#             raise Exception("GitHub token expired or invalid. Please re-login via frontend.")
#         ref_url = f"https://api.github.com/repos/{username}/{reponame}/git/ref/heads/{base_branch}"
#         ref_data = requests.get(ref_url, headers=self._github_headers()).json()
#         base_sha = ref_data["object"]["sha"]

#         create_resp = requests.post(
#             f"https://api.github.com/repos/{username}/{reponame}/git/refs",
#             headers=self._github_headers(),
#             json={"ref": f"refs/heads/{feature_branch}", "sha": base_sha},
#         )

#         if create_resp.status_code == 201:
#             print(f"✅ [ACTION] Created new branch: {feature_branch}")
#         elif create_resp.status_code == 422:
#             print(f"⚠️ [ACTION] Branch {feature_branch} already exists (reusing).")
#         elif create_resp.status_code == 401:
#             print("❌ [AUTH] Invalid GitHub token during branch creation.")
#             raise Exception("GitHub token expired. Please re-authenticate.")
#         else:
#             print(f"❌ [ACTION] Failed to create branch: {create_resp.text}")

#     def _get_file_content_and_sha(self, username, reponame, branch, path):
#         resp = requests.get(
#             f"https://api.github.com/repos/{username}/{reponame}/contents/{path}?ref={branch}",
#             headers=self._github_headers(),
#         )
#         resp.raise_for_status()
#         data = resp.json()
#         content = base64.b64decode(data["content"]).decode("utf-8", errors="ignore")
#         return content, data["sha"]

#     def _commit_file_update(self, username, reponame, branch, path, new_content, commit_message, sha):
#         if not self._validate_github_token():
#             raise Exception("GitHub token expired or invalid. Please re-login via frontend.")

#         encoded = base64.b64encode(new_content.encode("utf-8")).decode("utf-8")
#         resp = requests.put(
#             f"https://api.github.com/repos/{username}/{reponame}/contents/{path}",
#             headers=self._github_headers(),
#             json={"message": commit_message, "content": encoded, "branch": branch, "sha": sha},
#         )

#         if resp.status_code == 201 or resp.status_code == 200:
#             print(f"✅ [ACTION] Committed {path} → {branch}")
#         elif resp.status_code == 401:
#             print("❌ [AUTH] Invalid GitHub token during commit.")
#             raise Exception("GitHub token expired or invalid. Please re-login.")
#         else:
#             raise Exception(f"Commit failed: {resp.text}")



#     def apply_suggestions(self, file_path, original_content, suggestions):
#         """
#         Use Gemini to apply accepted suggestions directly to the code.
#         The AI will rewrite the code implementing the changes intelligently.
#         """
#         print(f"\n[AI-FIX] ✏️ Applying {len(suggestions)} suggestion(s) to {file_path}")

#         # Concatenate all developer-approved comments into one instruction block
#         suggestions_text = "\n".join(
#             [f"- {s.get('comment', '')}" for s in suggestions if s.get("comment")]
#         )

#         ext = os.path.splitext(file_path)[1]

#         prompt = f"""
#         You are an AI software engineer.
#         Modify the following code based on the accepted refactor suggestions.
#         Return only the updated code (no explanations, no markdown).
    
#         File type: {ext}
#         Accepted suggestions:
#         {suggestions_text}
    
#         Original code:
#         {original_content}
#         """

#         try:
#             result = self.model.generate_content(prompt)
#             updated_code = result.text.strip()

#             if not updated_code:
#                 print(f"⚠️ [AI-FIX] No updated code returned for {file_path}. Keeping original.")
#                 return original_content

#             print(f"✅ [AI-FIX] Successfully updated {file_path}")
#             return updated_code
#         except Exception as e:
#             print(f"❌ [AI-FIX] Error applying suggestions to {file_path}: {e}")
#             return original_content


#     # ----------------------------------------
#     # 4️⃣ SMART PROJECT-LEVEL DOCS
#     # ----------------------------------------
#     def generate_project_docs(self, repo_name: str, files: list, analyses: list, format: str = "md") -> str:
#         """
#         Generates a structured Technical Design Document (TDD) using an LLM 
#         based on the provided repository analysis and project context.
    
#         Args:
#             repo_name: The name of the GitHub repository/project.
#             files: List of all files in the repository.
#             analyses: List of analysis summaries generated by the AI for each file.
#             format: The desired output format (default is "md").
    
#         Returns:
#             The generated TDD documentation in Markdown format.
#         """
#         print(f"\n[PROJECT DOCS] 🏗️ Generating {format.upper()} TDD documentation for: {repo_name}")
    
#         # Concatenate file summaries into a single context string
#         summarized = "\n".join(
#             [f"- {a.get('file', 'N/A')}: {a.get('summary', '')}" for a in analyses if a.get("summary")]
#         )

#         project_prompt = f"""
#     You are a senior software architect and documentation expert.
#     Based on the following repository and analysis data, generate a **Technical Design Document (TDD)** for the project named **{repo_name}**.

#     ---
#     ### Document Characteristics
#     - **Document Type:** PROJECT DOCUMENTATION (Technical Design Document)
#     - **Audience:** Technical stakeholders — AI/ML Team, Developers, Engineering Managers
#     - **Purpose:** To translate the project into a detailed, actionable engineering plan
#     - **Tone & Style:** Formal, definitive, and structured (hierarchical numbering like 1.0, 1.1, 1.2)
#     - **Format:** Include Title Page, Table of Contents, and structured sections

#     ---
#     ### Standardized Section Template
#     The generated documentation **must** include the following sections in this order:

#     **1. Project Overview**
#     - 1.1 Project Summary
#     - 1.2 Problem Statement
#     - 1.3 Business Use Case
#     - 1.4 Target Use Case

#     **2. Functionality & Features**
#     - 2.1 Multi-Format File Support
#     - 2.2 Core AI or ML Features
#     - 2.3 Smart Automation or Optimization Components

#     **3. Architecture & Workflow**
#     - 3.1 High-Level Architecture (textual description)
#     - 3.2 Execution Flow (logical steps of core process)
#     - 3.3 Module Responsibilities

#     **4. Technology Stack**
#     - 4.1 Programming Language(s)
#     - 4.2 Libraries and Frameworks Used
#     - 4.3 External Tools or APIs

#     **5. Setup Instructions**
#     - 5.1 System Requirements
#     - 5.2 Installation Guide
#     - 5.3 Run Instructions

#     **6. Deployment & Testing**
#     - 6.1 Deployment Strategy
#     - 6.2 Testing & Quality Assurance Plan

#     **7. Appendix**
#     - 7.1 Glossary of Key Terms and Acronyms

#     ---
#     ### Repository Context
#     Below is the project context derived from the analysis:

#     {summarized}

#     Generate the document in a clean, professional Markdown format.
#     Ensure each section is **visually separated** with horizontal lines and proper spacing.
#     Include a **Title Page** at the top:

#     ```
#     --------------------------------------------------------
#                           PROJECT DOCUMENTATION
#                     Technical Design Document (TDD)
#                              Project: {repo_name}
#     --------------------------------------------------------
#     ```
#     """

#         try:
#             # Assuming self.model is the configured Gemini client
#             res = self.model.generate_content(project_prompt)
#             final_doc = res.text.strip()

#             # Add page-like border styling (for PDF export)
#             bordered_doc = f"""


#     🧠 TECHNICAL DESIGN DOCUMENT (TDD)
#     Project: {repo_name}


#     {final_doc}


#     Generated by CodeZen AI Developer Agent
#     """
#             print("✅ [PROJECT DOCS] TDD-style documentation generated successfully.")
#             return bordered_doc
#         except Exception as e:
#             print(f"❌ [PROJECT DOCS] Error: {e}")
#             return f"Error generating TDD-style project docs: {e}"
#     # ----------------------------------------
#     # 5️⃣ LEARNING — Store user feedback locally
#     # ----------------------------------------
#     def update_user_preferences(self, user_id, accepted, rejected):
#         memory_file = ".codezen_memory.json"
#         if os.path.exists(memory_file):
#             try:
#                 with open(memory_file, "r", encoding="utf-8") as f:
#                     memory_data = json.load(f)
#             except Exception:
#                 memory_data = {}
#         else:
#             memory_data = {}

#         if user_id not in memory_data:
#             memory_data[user_id] = {"accepted": [], "rejected": []}

#         accepted_data = [{"file_name": s.get("file_name"), "comment": s.get("comment"), "timestamp": time.time()} for s in accepted or []]
#         rejected_data = [{"file_name": s.get("file_name"), "comment": s.get("comment"), "timestamp": time.time()} for s in rejected or []]

#         memory_data[user_id]["accepted"].extend(accepted_data)
#         memory_data[user_id]["rejected"].extend(rejected_data)

#         with open(memory_file, "w", encoding="utf-8") as f:
#             json.dump(memory_data, f, indent=2)

#         self.user_preferences[user_id] = memory_data[user_id]
#         print(f"🧠 [LEARNING] Updated user preferences for '{user_id}' — {len(accepted_data)} accepted, {len(rejected_data)} rejected suggestions.")
        
#     def apply_changes_and_generate_docs(
#         self,
#         repo_url,
#         user_id="default_user",
#         accepted_suggestions=None,
#         rejected_suggestions=None,
#         doc_format="md",
#         commit_changes=False,
#     ):
#         """Alias wrapper for backward compatibility."""
#         return self.run_full_pipeline(
#             repo_url,
#             user_id,
#             accepted_suggestions,
#             rejected_suggestions,
#             doc_format,
#             commit_changes,
#         )

#     # ----------------------------------------
#     # 7️⃣ OPTIONAL — Auto Create Pull Request after AI commits
#     # ----------------------------------------
#     def _create_pull_request(self, username, reponame, feature_branch, base_branch):
#         """Automatically open a pull request on GitHub after AI commits."""
#         try:
#             pr_url = f"https://api.github.com/repos/{username}/{reponame}/pulls"
#             headers = self._github_headers()
#             pr_title = f"🤖 CodeZen AI Refactors — {feature_branch}"
#             pr_body = (
#                 f"### Automated Code Improvements by CodeZen\n\n"
#                 f"CodeZen AI applied developer-approved refactors and improvements.\n\n"
#                 f"**Base Branch:** `{base_branch}`\n"
#                 f"**Feature Branch:** `{feature_branch}`\n\n"
#                 f"Please review the changes and merge if they look good. 🚀"
#             )

#             payload = {
#                 "title": pr_title,
#                 "head": feature_branch,
#                 "base": base_branch,
#                 "body": pr_body
#             }

#             response = requests.post(pr_url, headers=headers, json=payload)

#             if response.status_code in [200, 201]:
#                 pr_data = response.json()
#                 print(f"✅ [PR] Pull Request Created: {pr_data.get('html_url')}")
#                 return pr_data.get("html_url")
#             elif response.status_code == 422:
#                 print("⚠️ [PR] A pull request for this branch already exists.")
#             elif response.status_code == 401:
#                 print("❌ [PR] Unauthorized — GitHub token invalid or missing.")
#             else:
#                 print(f"❌ [PR] Failed to create pull request: {response.text}")

#         except Exception as e:
#             print(f"❌ [PR] Error creating pull request: {e}")

#     # ----------------------------------------
#     # 6️⃣ FULL PIPELINE — Per-file analysis + unified doc + commit
#     # ----------------------------------------
#     # def run_full_pipeline(
#     #     self,
#     #     repo_url,
#     #     user_id="default_user",
#     #     accepted_suggestions=None,
#     #     rejected_suggestions=None,
#     #     doc_format="md",
#     #     commit_changes=False,
#     # ):
#     #     print(f"\n🚀 [INIT] Running AI Developer pipeline for: {repo_url}")

#     #     # 🧠 PERCEPTION — Fetch repository files
#     #     repo = self.fetch_repo(repo_url)
#     #     files = repo["files"]
#     #     if not files:
#     #         return {
#     #             "repo": repo["repo_name"],
#     #             "analysis": [],
#     #             "project_doc": "No code files found.",
#     #         }

#     #     # 🧩 REASONING — Analyze each file
#     #     analyses = []
#     #     for f in files:
#     #         a = self.analyze_code(f["file_name"], f["content"])
#     #         a["file"] = f["file_name"]
#     #         analyses.append(a)

#     #     # ⚙️ ACTION — Apply and commit accepted suggestions (optional)
#     #     if commit_changes and accepted_suggestions:
#     #         print("\n[ACTION] 🧩 Applying developer-approved suggestions and committing changes...")
#     #         username, reponame = self._parse_repo(repo_url)
#     #         base_branch = self._get_default_branch(username, reponame)
#     #         feature_branch = f"codezen/fixes-{int(time.time())}"
#     #         self._ensure_feature_branch(username, reponame, base_branch, feature_branch)

#     #         grouped = {}

#     #         # ✅ Handles both dicts and pydantic models
#     #         for s in accepted_suggestions:
#     #             accepted = s.get("accepted") if isinstance(s, dict) else getattr(s, "accepted", False)
#     #             file_name = s.get("file_name") if isinstance(s, dict) else getattr(s, "file_name", None)
#     #             comment = s.get("comment") if isinstance(s, dict) else getattr(s, "comment", "")

#     #             if accepted and file_name:
#     #                 grouped.setdefault(file_name, []).append({
#     #                     "file_name": file_name,
#     #                     "comment": comment,
#     #                     "accepted": accepted
#     #                 })

#     #         # ✅ Skip .ipynb corruption and only modify safe files
#     #         for file_name, suggs in grouped.items():
#     #             if file_name.endswith(".ipynb"):
#     #                 print(f"⚠️ [NOTEBOOK] Skipping direct modification for {file_name} to prevent corruption.")
#     #                 continue
#     #             try:
#     #                 current_content, sha = self._get_file_content_and_sha(username, reponame, feature_branch, file_name)
#     #                 updated_content = self.apply_suggestions(file_name, current_content, suggs)
#     #                 message = f"AI-Fix: Applied {len(suggs)} developer-approved suggestion(s) to {file_name}"
#     #                 self._commit_file_update(username, reponame, feature_branch, file_name, updated_content, message, sha)
#     #             except Exception as e:
#     #                 print(f"❌ [ACTION] Failed to commit {file_name}: {e}")

#     #         print(f"✅ [ACTION] All accepted suggestions committed to branch: {feature_branch}")

#     #         # 🪄 OPTIONAL NEXT STEP (coming soon): Auto create pull request
#     #         self._create_pull_request(username, reponame, feature_branch, base_branch)
#     #     else:
#     #         print("[ACTION] ⚪ Skipping commit — No accepted suggestions or commit disabled.")

#     #     # 🧾 LEARNING — Update user memory
#     #     self.update_user_preferences(user_id, accepted_suggestions or [], rejected_suggestions or [])

#     #     # 📘 DOCUMENTATION — Generate project-level docs
#     #     project_doc = self.generate_project_docs(repo["repo_name"], files, analyses, format=doc_format)

#     #     print("✅ [COMPLETE] AI Developer pipeline executed successfully.")
#     #     return {
#     #         "repo": repo["repo_name"],
#     #         "analysis": analyses,
#     #         "project_doc": project_doc,
#     #         "user_learning": self.user_preferences.get(user_id, {}),
#     #     }
    
#     # ----------------------------------------

# # 6️⃣ FULL PIPELINE — Per-file analysis + unified doc + commit (Enhanced)

# # ----------------------------------------

# def run_full_pipeline(
#     self,
#     repo_url,
#     user_id="default_user",
#     accepted_suggestions=None,
#     rejected_suggestions=None,
#     doc_format="md",
#     commit_changes=False,
# ):
#     import difflib

#     ```
#     print(f"\n🚀 [INIT] Running AI Developer pipeline for: {repo_url}")

#     # ✅ Initialize per-user cache
#     if not hasattr(self, "session_memory"):
#         self.session_memory = {}
#     if user_id not in self.session_memory:
#         self.session_memory[user_id] = {}

#     # ⚡ Reuse cached analysis if available (for resume after OAuth)
#     if repo_url in self.session_memory[user_id]:
#         print("⚡ [CACHE] Using cached repository analysis for user:", user_id)
#         files, analyses = self.session_memory[user_id][repo_url]
#     else:
#         # 🧠 PERCEPTION — Fetch repository files
#         repo = self.fetch_repo(repo_url)
#         files = repo["files"]
#         if not files:
#             return {
#                 "repo": repo["repo_name"],
#                 "analysis": [],
#                 "project_doc": "No code files found.",
#             }

#         # 🧩 REASONING — Analyze each file
#         analyses = []
#         for f in files:
#             a = self.analyze_code(f["file_name"], f["content"])
#             a["file"] = f["file_name"]
#             analyses.append(a)

#         # Cache the analysis for this repo and user
#         self.session_memory[user_id][repo_url] = (files, analyses)

#     # ⚙️ ACTION — Apply and commit accepted suggestions (optional)
#     if commit_changes and accepted_suggestions:
#         print("\n[ACTION] 🧩 Applying developer-approved suggestions and committing changes...")
#         username, reponame = self._parse_repo(repo_url)
#         base_branch = self._get_default_branch(username, reponame)
#         feature_branch = f"codezen/fixes-{int(time.time())}"
#         self._ensure_feature_branch(username, reponame, base_branch, feature_branch)

#         grouped = {}

#         # ✅ Group accepted suggestions by file
#         for s in accepted_suggestions:
#             accepted = s.get("accepted") if isinstance(s, dict) else getattr(s, "accepted", False)
#             file_name = s.get("file_name") if isinstance(s, dict) else getattr(s, "file_name", None)
#             comment = s.get("comment") if isinstance(s, dict) else getattr(s, "comment", "")
#             if accepted and file_name:
#                 grouped.setdefault(file_name, []).append({
#                     "file_name": file_name,
#                     "comment": comment,
#                     "accepted": accepted
#                 })

#         # ✅ Process file changes
#         for file_name, suggs in grouped.items():
#             if file_name.endswith(".ipynb"):
#                 print(f"⚠️ [NOTEBOOK] Skipping direct modification for {file_name} to prevent corruption.")
#                 continue
#         try:
#             current_content, sha = self._get_file_content_and_sha(username, reponame, feature_branch, file_name)
#             updated_content = self.apply_suggestions(file_name, current_content, suggs)

#             # 🧾 Compute diff summary (what changed)
#             diff = difflib.unified_diff(
#                    current_content.splitlines(),
#                    updated_content.splitlines(),
#                    fromfile=f"{file_name} (original)",
#                    tofile=f"{file_name} (updated)",
#                    lineterm=""
#             )
#             diff_text = "\n".join(diff)

#             # Attach diff info to analysis for frontend preview
#             entry = next((a for a in analyses if a["file"] == file_name), None)
#             if entry:
#                 entry["diff"] = diff_text

#             message = f"AI-Fix: Applied {len(suggs)} developer-approved suggestion(s) to {file_name}"
#             self._commit_file_update(username, reponame, feature_branch, file_name, updated_content, message, sha)

#         except Exception as e:
#             print(f"❌ [ACTION] Failed to commit {file_name}: {e}")

#         print(f"✅ [ACTION] All accepted suggestions committed to branch: {feature_branch}")
#         self._create_pull_request(username, reponame, feature_branch, base_branch)

#     else:
#         print("[ACTION] ⚪ Skipping commit — No accepted suggestions or commit disabled.")

#     # 🧾 LEARNING — Update user memory
#     self.update_user_preferences(user_id, accepted_suggestions or [], rejected_suggestions or [])

#     # 📘 DOCUMENTATION — Generate project-level docs
#     repo_name = repo_url.split("/")[-1].replace(".git", "")
#     project_doc = self.generate_project_docs(repo_name, files, analyses, format=doc_format)

#     print("✅ [COMPLETE] AI Developer pipeline executed successfully.")
#     return {
#         "repo": repo_name,
#         "analysis": analyses,
#         "project_doc": project_doc,
#         "user_learning": self.user_preferences.get(user_id, {}),
#     }
    

# import json
# import time
# import requests
# import os
# import re
# import base64
# import difflib # <-- ADDED MISSING IMPORT
# import markdown
# from fpdf import FPDF
# from docx import Document
# from tempfile import NamedTemporaryFile
# import weasyprint
# from typing import Dict, Any, List, Optional
# from autogen import ConversableAgent
# import google.generativeai as genai
# from dotenv import load_dotenv
# load_dotenv()
# genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# # ----------------------------------------
# # Comment Syntax Configurations
# # ----------------------------------------
# COMMENT_PREFIX = {
#     ".py": "# ", ".js": "// ", ".java": "// ", ".c": "// ", ".cpp": "// ",
#     ".ts": "// ", ".html": "", ".css": " */"}


# class DeveloperAgent(ConversableAgent):
#     """
#     Unified AI Developer Agent for CodeZen. Handles GitHub interaction,
#     Gemini-based code analysis, refactoring, and documentation generation.
#     """
#     def __init__(self, model_name: str = "gemini-2.5-flash"):
#         # 🔵 STEP 1: UPDATED INIT AND TOKEN HANDLING METHODS
#         super().__init__(name="DeveloperAgent", description="Unified AI Developer Agent for CodeZen.")
#         self.model_name = model_name
#         self.model = genai.GenerativeModel(model_name)
#         self.user_preferences: Dict[str, Any] = {}
#         self.github_token: Optional[str] = None  # ✅ stores user token dynamically
#         self.session_memory: Dict[str, Dict[str, Any]] = {} # Cache for analysis
#         print(f"✅ DeveloperAgent initialized with {model_name}")

#     def set_dynamic_token(self, token: str):
#         """Store GitHub token from frontend OAuth for current user session."""
#         self.github_token = token
#         print("🔐 Dynamic GitHub token set successfully.")

#     # ----------------------------------------
#     # GitHub header helper (with token check)
#     # ----------------------------------------
#     def _github_headers(self) -> Dict[str, str]:
#         """Use dynamic token from OAuth if available, otherwise fall back to env variable."""
#         # Use getattr safely and fall back to os.getenv
#         token = self.github_token or os.getenv("GITHUB_TOKEN")
#         if not token:
#             raise Exception("Missing GitHub token (neither dynamic nor environment).")
#         return {"Authorization": f"token {token}", "Accept": "application/vnd.github+json"}

#     # ----------------------------------------
#     # Token validation helper (NEW)
#     # ----------------------------------------
#     def _validate_github_token(self) -> bool:
#         """Validate current GitHub token; returns True if valid, False if expired/invalid."""
#         try:
#             headers = self._github_headers()
#             response = requests.get("https://api.github.com/user", headers=headers)
#             if response.status_code == 200:
#                 return True
#             if response.status_code in [401, 403]:
#                 print("⚠️ [AUTH] GitHub token is invalid or expired.")
#                 self.github_token = None
#                 return False
#             print(f"⚠️ [AUTH] Unexpected GitHub API status: {response.status_code}")
#             return False
#         except Exception as e:
#             print(f"❌ [AUTH] Error validating token: {e}")
#             return False

#     # ----------------------------------------
#     # 1️⃣ PERCEPTION — Fetch repo recursively
#     # ----------------------------------------
#     def fetch_repo(self, repo_url: str) -> Dict[str, Any]:
#         print(f"\n[PERCEPTION] 🌐 Fetching repository structure from: {repo_url}")
#         parts = repo_url.strip("/").split("/")
#         # The user's repo URL structure seems to be /username/reponame
#         if len(parts) < 2:
#              raise ValueError("Invalid GitHub repository URL format.")
             
#         username, reponame = parts[-2], parts[-1].replace(".git", "")
        
#         # Need headers for this call to work reliably, especially for private repos
#         headers = self._github_headers() if self.github_token or os.getenv("GITHUB_TOKEN") else {}
        
#         meta_resp = requests.get(f"https://api.github.com/repos/{username}/{reponame}", headers=headers)
#         meta_resp.raise_for_status() # Raise for bad status codes
#         meta = meta_resp.json()
#         default_branch = meta.get("default_branch", "main")
#         print(f"[PERCEPTION] 🪶 Default branch: {default_branch}")

#         tree_resp = requests.get(
#             f"https://api.github.com/repos/{username}/{reponame}/git/trees/{default_branch}?recursive=1",
#             headers=headers
#         )
#         tree_resp.raise_for_status()
#         tree = tree_resp.json()
        
#         files: List[Dict[str, str]] = []
#         allowed = [".py", ".js", ".java", ".ts", ".cpp", ".c", ".html", ".css", ".ipynb"]
        
#         for item in tree.get("tree", []):
#             if item["type"] == "blob" and any(item["path"].endswith(ext) for ext in allowed):
#                 raw = f"https://raw.githubusercontent.com/{username}/{reponame}/{default_branch}/{item['path']}"
                
#                 # Use stream=True and get content using raw URL for large files, though
#                 # GitHub API content endpoint (used in _get_file_content_and_sha) is usually better.
#                 # Sticking to the raw URL for simplicity as per original.
#                 content_resp = requests.get(raw)
#                 content_resp.raise_for_status()
#                 content = content_resp.text
                
#                 if item["path"].endswith(".ipynb"):
#                     try:
#                         nb = json.loads(content)
#                         cells = [
#                             "\n".join(c.get("source", []))
#                             for c in nb.get("cells", [])
#                             if c.get("cell_type") == "code"
#                         ]
#                         content = "\n\n".join(cells)
#                         print(f"📘 [NOTEBOOK] Extracted {len(cells)} code cells from {item['path']}")
#                     except Exception as e:
#                         print(f"⚠️ [NOTEBOOK] Failed to parse notebook {item['path']}: {e}")
#                         # Fallback: Treat as a regular text file if parsing fails
                        
#                 files.append({"file_name": item["path"], "content": content})
                
#         print(f"✅ [PERCEPTION] Repo fetched — {len(files)} code files detected.")
#         return {"repo_name": reponame, "files": files}

#     # ----------------------------------------
#     # 2️⃣ REASONING — Analyze each file with Gemini
#     # ----------------------------------------
#     def analyze_code(self, file_name: str, content: str) -> Dict[str, Any]:
#         print(f"\n[REASONING] 🧩 Analyzing file: {file_name}")
#         prompt = f"""
#         You are an expert AI code reviewer.
#         Analyze the following code and return **ONLY valid JSON** with no extra text or markdown.

#         The JSON object MUST follow this structure:
#         {{
#           "issues": ["list of problems or bugs (e.g., hardcoded secrets, security risks)"],
#           "refactors": ["list of improvement suggestions (e.g., better structure, modularity)"],
#           "summary": "short summary of what this file does"
#         }}
        
#         Code:
#         ```
#         {content}
#         ```
#         """
#         try:
#             # Use JSON mode if supported by the model and API
#             # For simplicity, keeping the original logic of string parsing.
#             out = self.model.generate_content(prompt).text.strip()
            
#             try:
#                 parsed = json.loads(out)
#             except json.JSONDecodeError:
#                 # Attempt to extract JSON from surrounding text
#                 match = re.search(r"\{.*\}", out, re.DOTALL)
#                 if match:
#                     # Clean up common issues before final parse
#                     clean_json_str = match.group().replace('```json', '').replace('```', '').strip()
#                     parsed = json.loads(clean_json_str)
#                 else:
#                     # Fallback to a basic structure if no JSON found
#                     parsed = {"issues": [], "refactors": [], "summary": out[:300]}
            
#             print(f"✅ [REASONING] Parsed for {file_name}")
#             return parsed
            
#         except Exception as e:
#             print(f"❌ [REASONING] Error: {e}")
#             return {"file": file_name, "issues": [], "refactors": [], "summary": str(e)}

#     # ----------------------------------------
#     # 3️⃣ ACTION — GitHub Commit Helpers
#     # ----------------------------------------
#     def _parse_repo(self, repo_url: str) -> tuple[str, str]:
#         parts = repo_url.strip("/").split("/")
#         return parts[-2], parts[-1].replace(".git", "")

#     def _get_default_branch(self, username: str, reponame: str) -> str:
#         if not self._validate_github_token():
#             raise Exception("GitHub token expired or invalid. Please re-login via frontend.")
#         resp = requests.get(f"https://api.github.com/repos/{username}/{reponame}", headers=self._github_headers())
#         resp.raise_for_status()
#         return resp.json().get("default_branch", "main")

#     def _ensure_feature_branch(self, username: str, reponame: str, base_branch: str, feature_branch: str):
#         if not self._validate_github_token():
#             raise Exception("GitHub token expired or invalid. Please re-login via frontend.")
            
#         # 1. Get SHA of the base branch
#         ref_url = f"https://api.github.com/repos/{username}/{reponame}/git/ref/heads/{base_branch}"
#         ref_resp = requests.get(ref_url, headers=self._github_headers())
#         ref_resp.raise_for_status()
#         ref_data = ref_resp.json()
#         base_sha = ref_data["object"]["sha"]

#         # 2. Try to create the new feature branch
#         create_resp = requests.post(
#             f"https://api.github.com/repos/{username}/{reponame}/git/refs",
#             headers=self._github_headers(),
#             json={"ref": f"refs/heads/{feature_branch}", "sha": base_sha},
#         )

#         if create_resp.status_code == 201:
#             print(f"✅ [ACTION] Created new branch: {feature_branch}")
#         elif create_resp.status_code == 422 and "Reference already exists" in create_resp.text:
#             print(f"⚠️ [ACTION] Branch {feature_branch} already exists (reusing).")
#         elif create_resp.status_code == 401:
#             print("❌ [AUTH] Invalid GitHub token during branch creation.")
#             raise Exception("GitHub token expired. Please re-authenticate.")
#         else:
#             print(f"❌ [ACTION] Failed to create branch: {create_resp.text}")
#             create_resp.raise_for_status() # Raise for any other error

#     def _get_file_content_and_sha(self, username: str, reponame: str, branch: str, path: str) -> tuple[str, str]:
#         resp = requests.get(
#             f"https://api.github.com/repos/{username}/{reponame}/contents/{path}?ref={branch}",
#             headers=self._github_headers(),
#         )
#         resp.raise_for_status()
#         data = resp.json()
#         content = base64.b64decode(data["content"]).decode("utf-8", errors="ignore")
#         return content, data["sha"]

#     def _commit_file_update(self, username: str, reponame: str, branch: str, path: str, new_content: str, commit_message: str, sha: str):
#         if not self._validate_github_token():
#             raise Exception("GitHub token expired or invalid. Please re-login via frontend.")

#         encoded = base64.b64encode(new_content.encode("utf-8")).decode("utf-8")
#         resp = requests.put(
#             f"https://api.github.com/repos/{username}/{reponame}/contents/{path}",
#             headers=self._github_headers(),
#             json={"message": commit_message, "content": encoded, "branch": branch, "sha": sha},
#         )

#         if resp.status_code in [200, 201]:
#             print(f"✅ [ACTION] Committed {path} → {branch}")
#         elif resp.status_code == 401:
#             print("❌ [AUTH] Invalid GitHub token during commit.")
#             raise Exception("GitHub token expired or invalid. Please re-login.")
#         else:
#             resp.raise_for_status() # Raise for any other error

#     def apply_suggestions(self, file_path: str, original_content: str, suggestions: List[Dict[str, Any]]) -> str:
#         """
#         Use Gemini to apply accepted suggestions directly to the code by rewriting it.
#         """
#         print(f"\n[AI-FIX] ✏️ Applying {len(suggestions)} suggestion(s) to {file_path}")

#         # Concatenate all developer-approved comments into one instruction block
#         suggestions_text = "\n".join(
#             [f"- {s.get('comment', '')}" for s in suggestions if s.get("comment")]
#         )

#         ext = os.path.splitext(file_path)[1]

#         prompt = f"""
#         You are an AI software engineer.
#         Modify the following code based on the accepted refactor suggestions.
#         Return **ONLY the updated code** (do not include explanations, comments, or markdown formatting like ```).
    
#         File type: {ext}
#         Accepted suggestions:
#         {suggestions_text}
    
#         Original code:
#         {original_content}
#         """

#         try:
#             result = self.model.generate_content(prompt)
#             updated_code = result.text.strip()

#             if not updated_code:
#                 print(f"⚠️ [AI-FIX] No updated code returned for {file_path}. Keeping original.")
#                 return original_content

#             print(f"✅ [AI-FIX] Successfully updated {file_path}")
#             return updated_code
#         except Exception as e:
#             print(f"❌ [AI-FIX] Error applying suggestions to {file_path}: {e}")
#             return original_content
    
    
#     # ----------------------------------------
#     # 4️⃣ SMART PROJECT-LEVEL DOCS (Enhanced)
#     # ----------------------------------------
   
#     def convert_doc(markdown_text: str, output_format: str = "pdf") -> str:
#         """
#         Convert Markdown text into either a PDF or DOCX file.
#         Returns the path to the generated file.
#         """
#         output_format = output_format.lower().strip()
#         file = NamedTemporaryFile(delete=False, suffix=f".{output_format}")
#         output_path = file.name

#         if output_format == "pdf":
#             html = markdown.markdown(markdown_text)
#             try:
#                 weasyprint.HTML(string=html).write_pdf(output_path)
#             except Exception:
#                 pdf = FPDF()
#                 pdf.add_page()
#                 pdf.set_font("Arial", size=12)
#                 for line in html.splitlines():
#                     pdf.multi_cell(0, 10, line)
#                 pdf.output(output_path)

#         elif output_format == "docx":
#             doc = Document()
#             for line in markdown_text.split("\n"):
#                 doc.add_paragraph(line)
#             doc.save(output_path)

#         else:
#             raise ValueError("Unsupported format. Choose 'pdf' or 'docx'.")

#         return output_path

#     def generate_project_docs(
#         self,
#         repo_name: str,
#         files: List[Dict[str, str]],
#         analyses: List[Dict[str, Any]],
#         format: str = "pdf"
#     ) -> str:
#         """
#         Generates a README, TDD, or structured document.
#         Supports export as Markdown, PDF, or DOCX.
#         """
#         print(f"\n[PROJECT DOCS] 🏗️ Generating {format.upper()} documentation for: {repo_name}")

#         # Concatenate summaries
#         summarized = "\n".join(
#             [f"- **{a.get('file', 'N/A')}** → {a.get('summary', '')}" for a in analyses if a.get("summary")]
#         )

#         # Template for detailed TDD
#         project_prompt = f"""
#         You are a senior software architect and documentation expert.
#         Based on the following repository and analysis data, generate a **Technical Design Document (TDD)** 
#         for the project named **{repo_name}**.

#         ---
#         ### Document Characteristics
#         - **Type:** Technical Design Document (TDD)
#         - **Audience:** Technical stakeholders — AI/ML Team, Developers, Engineering Managers
#         - **Purpose:** Explain system design, architecture, and engineering logic
#         - **Tone:** Formal, structured, and explanatory
#         - **Format:** Markdown with hierarchical numbering (1.0, 1.1, etc.)
#         ---

#         ### Standardized Section Template

#         1. Project Overview  
#             1.1 Project Summary  
#             1.2 Problem Statement  
#             1.3 Business Use Case  
#             1.4 Target Users  

#         2. Functionality & Features  
#             2.1 Core Functionalities  
#             2.2 File or Data Flow Description  

#         3. Architecture & Workflow  
#             3.1 High-Level Architecture  
#             3.2 Module Responsibilities  
#             3.3 Execution Flow  
#    -        Include a simple ASCII diagram like:
#             ```
#             [Input Directory] → [main.py] → [language_detector.py] → [Output Directory]
#             ```

#         4. Technology Stack  
#             4.1 Programming Languages  
#             4.2 Libraries & Tools  
#             4.3 External Services or APIs  

#         5. Setup Instructions  
#             5.1 System Requirements  
#             5.2 Installation & Configuration  
#             5.3 How to Run the Application  

#         6. Testing & Validation  
#             - Describe how each module can be tested.
#             - Include a table like:
#             | Test Case | Input | Expected Output | Status |
#             |------------|--------|----------------|---------|
#             | Valid Input | sample.mp3 | Detected Language = English | ✅ |
#             | Invalid File | corrupt.wav | Error Logged | ⚠️ |

#         7. Prompt Engineering & Strategy (for LLM-based projects)  
#             7.1 Prompt Design Principles  
#             7.2 Example Prompts  
#             7.3 Weighting and Evaluation Logic  

#         8. Troubleshooting  
#             8.1 Common Issues  
#             8.2 Possible Causes and Solutions  

#         9. Future Scope  
#             9.1 Feature Enhancements  
#             9.2 Scalability and Deployment Ideas  

#         10. Appendix  
#             10.1 Glossary  
#             10.2 References  

#         ---
#         ### Repository Context
#         {summarized}

#         Generate the document in Markdown format, formatted with clear numbering and subheadings.
#         """

#             # Generate with Gemini
#             try:
#                 res = self.model.generate_content(project_prompt)
#                 final_doc = res.text.strip()
#                 print("✅ [PROJECT DOCS] Raw documentation generated successfully.")

#                 # Ensure essential sections exist
#                 if "Prompt Engineering" not in final_doc:
#                     final_doc += """
#         ## 7. Prompt Engineering & Strategy
#         Describe how the LLM or AI component interacts with the project.
#         Include input-output prompt examples, decision logic, and tone control.
#         """

#                 if "Troubleshooting" not in final_doc:
#                     final_doc += """
#         ## 8. Troubleshooting
#         | Problem | Possible Cause | Solution |
#         |----------|----------------|-----------|
#         | LLM output incomplete | Prompt too long | Chunk or summarize input |
#         | API request failed | Missing token | Check environment variables |
#         """

#                 if "Future Scope" not in final_doc:
#                     final_doc += """
#         ## 9. Future Scope
#         - Add containerization with Docker.
#         - Deploy on cloud for scalability.
#         - Integrate database or analytics pipeline.
#         """

#             # Convert to desired output
#                 if format.lower() in ["pdf", "docx"]:
#                     print(f"📄 [CONVERT] Converting to {format.upper()} format...")
#                     output_path = convert_doc(final_doc, output_format=format)
#                     print(f"✅ [PROJECT DOCS] Saved as {output_path}")
#                     return output_path
#                 else:
#                     bordered_doc = f"\n---\nPROJECT DOCUMENTATION: {repo_name}\n---\n\n{final_doc}\n\nGenerated by CodeZen AI Developer Agent"
#                 return bordered_doc

#             except Exception as e:
#                 print(f"❌ [PROJECT DOCS] Error: {e}")
#                 return f"Error generating documentation: {e}"


#     # # ----------------------------------------
#     # # 4️⃣ SMART PROJECT-LEVEL DOCS
#     # # ----------------------------------------
#     # def generate_project_docs(
#     #     self,
#     #     repo_name: str,
#     #     files: List[Dict[str, str]],
#     #     analyses: List[Dict[str, Any]],
#     #     format: str = "md"
#     # ) -> str:
#     #     """
#     #     Generates either a README.md or a full Technical Design Document (TDD)
#     #     based on the selected output format.
#     #     """
#     #     print(f"\n[PROJECT DOCS] 🏗️ Generating {format.upper()} documentation for: {repo_name}")


#     #     # Concatenate file summaries into a single context string
#     #     summarized = "\n".join(
#     #         [f"- {a.get('file', 'N/A')}: {a.get('summary', '')}" for a in analyses if a.get("summary")]
#     #     )

#     #     # Decide document type based on format
#     #     if format.lower() == "md":
#     #         # Generate README.md style summary
#     #         header = f"# {repo_name}\n\n"
#     #         project_prompt = f"""
#     #         You are a software documentation expert.
#     #         Create a **README.md** file for a GitHub repository named **{repo_name}**.
#     #         The README should summarize the project based on the file analyses below.

#     #         ### Requirements
#     #         - Keep it concise, professional, and Markdown formatted.
#     #         - Include sections like:
#     #         - Project Overview
#     #         - Features
#     #         - Tech Stack
#     #         - Setup & Usage
#     #         - Author / Credits (optional)
#     #         - The README should explain the purpose and usage of the project in simple terms.

#     #         --- Repository Context ---
#     #         {summarized}

#     #         Output only the README content in Markdown format.
#     #         """
#     #     else:
#     #         # Generate Technical Design Document (TDD)
#     #         header = f"---\nPROJECT DOCUMENTATION: {repo_name}\n---\n"
#     #         project_prompt = f"""
#     #         You are a senior software architect and documentation expert.
#     #         Based on the following repository and analysis data, generate a **Technical Design Document (TDD)** 
#     #         for the project named **{repo_name}**.

#     #         ---
#     #         ### Document Characteristics
#     #         - **Document Type:** PROJECT DOCUMENTATION (Technical Design Document)
#     #         - **Audience:** Technical stakeholders — AI/ML Team, Developers, Engineering Managers
#     #         - **Purpose:** To translate the project into a detailed, actionable engineering plan
#     #         - **Tone & Style:** Formal, definitive, and structured (hierarchical numbering like 1.0, 1.1, 1.2)
#     #         - **Format:** The final output MUST be in a clean Markdown format.

#     #         ---
#     #         ### Standardized Section Template
#     #         1. Project Overview
#     #         2. Functionality & Features
#     #         3. Architecture & Workflow
#     #         4. Technology Stack
#     #         5. Setup Instructions
#     #         6. Deployment & Testing
#     #         7. Appendix (Glossary / References)

#     #         ---
#     #         ### Repository Context
#     #         {summarized}

#     #         Generate the document in a professional Markdown format.
#     #         Include a clear title section at the top.
#     #         """

#     #     try:
#     #         res = self.model.generate_content(project_prompt)
#     #         final_doc = res.text.strip()

#     #         # Add header wrapper for clarity (for previews or exports)
#     #         bordered_doc = f"""

#     # {header}
#     # {final_doc}

#     # Generated by CodeZen AI Developer Agent
#     # """
#     #         print("✅ [PROJECT DOCS] Documentation generated successfully.")
#     #         return bordered_doc

#     #     except Exception as e:
#     #         print(f"❌ [PROJECT DOCS] Error: {e}")
#     #         return f"Error generating documentation: {e}"

            
#     # ----------------------------------------
#     # 5️⃣ LEARNING — Store user feedback locally
#     # ----------------------------------------
#     def update_user_preferences(self, user_id: str, accepted: Optional[List[Dict[str, Any]]], rejected: Optional[List[Dict[str, Any]]]):
#         memory_file = ".codezen_memory.json"
#         memory_data: Dict[str, Any] = {}
        
#         if os.path.exists(memory_file):
#             try:
#                 with open(memory_file, "r", encoding="utf-8") as f:
#                     memory_data = json.load(f)
#             except Exception:
#                 memory_data = {}

#         if user_id not in memory_data:
#             memory_data[user_id] = {"accepted": [], "rejected": []}

#         current_time = time.time()
        
#         # Use more explicit type checking and safety for accepted/rejected lists
#         accepted_data = [
#             {"file_name": s.get("file_name"), "comment": s.get("comment"), "timestamp": current_time} 
#             for s in accepted or [] if isinstance(s, dict)
#         ]
#         rejected_data = [
#             {"file_name": s.get("file_name"), "comment": s.get("comment"), "timestamp": current_time} 
#             for s in rejected or [] if isinstance(s, dict)
#         ]

#         memory_data[user_id]["accepted"].extend(accepted_data)
#         memory_data[user_id]["rejected"].extend(rejected_data)

#         try:
#             with open(memory_file, "w", encoding="utf-8") as f:
#                 json.dump(memory_data, f, indent=2)
#             self.user_preferences[user_id] = memory_data[user_id]
#             print(f"🧠 [LEARNING] Updated user preferences for '{user_id}' — {len(accepted_data)} accepted, {len(rejected_data)} rejected suggestions.")
#         except Exception as e:
#             print(f"❌ [LEARNING] Failed to write memory file: {e}")


#     def apply_changes_and_generate_docs(
#         self,
#         repo_url: str,
#         user_id: str = "default_user",
#         accepted_suggestions: Optional[List[Dict[str, Any]]] = None,
#         rejected_suggestions: Optional[List[Dict[str, Any]]] = None,
#         doc_format: str = "md",
#         commit_changes: bool = False,
#     ) -> Dict[str, Any]:
#         """Alias wrapper for backward compatibility."""
#         return self.run_full_pipeline(
#             repo_url,
#             user_id,
#             accepted_suggestions,
#             rejected_suggestions,
#             doc_format,
#             commit_changes,
#             use_cache=True,
#         )

#     # ----------------------------------------
#     # 7️⃣ OPTIONAL — Auto Create Pull Request after AI commits
#     # ----------------------------------------
#     def _create_pull_request(self, username: str, reponame: str, feature_branch: str, base_branch: str) -> Optional[str]:
#         """Automatically open a pull request on GitHub after AI commits."""
#         try:
#             # pr_url = f"[https://api.github.com/repos/](https://api.github.com/repos/){username}/{reponame}/pulls"
#             pr_url = f"https://api.github.com/repos/{username}/{reponame}/pulls"
#             headers = self._github_headers()
#             pr_title = f"🤖 CodeZen AI Refactors — {feature_branch}"
#             pr_body = (
#                 f"### Automated Code Improvements by CodeZen\n\n"
#                 f"CodeZen AI applied developer-approved refactors and improvements.\n\n"
#                 f"**Base Branch:** `{base_branch}`\n"
#                 f"**Feature Branch:** `{feature_branch}`\n\n"
#                 f"Please review the changes and merge if they look good. 🚀"
#             )

#             payload = {
#                 "title": pr_title,
#                 "head": feature_branch,
#                 "base": base_branch,
#                 "body": pr_body
#             }

#             response = requests.post(pr_url, headers=headers, json=payload)

#             if response.status_code in [200, 201]:
#                 pr_data = response.json()
#                 print(f"✅ [PR] Pull Request Created: {pr_data.get('html_url')}")
#                 return pr_data.get("html_url")
#             if response.status_code == 422:
#                 print("⚠️ [PR] A pull request for this branch already exists (or base/head error).")
#             elif response.status_code == 401:
#                 print("❌ [PR] Unauthorized — GitHub token invalid or missing.")
#             else:
#                 print(f"❌ [PR] Failed to create pull request: {response.text}")

#         except Exception as e:
#             print(f"❌ [PR] Error creating pull request: {e}")
#         return None

#     # ----------------------------------------
#     # 6️⃣ FULL PIPELINE — Perception, Reasoning, Action, and Documentation (Enhanced)
#     # ----------------------------------------
#     # def run_full_pipeline(
#     #     self,
#     #     repo_url: str,
#     #     user_id: str = "default_user",
#     #     accepted_suggestions: Optional[List[Dict[str, Any]]] = None,
#     #     rejected_suggestions: Optional[List[Dict[str, Any]]] = None,
#     #     doc_format: str = "md",
#     #     commit_changes: bool = False,
#     #     analysis_only: bool = False
#     # ) -> Dict[str, Any]:
#     #     """
#     #     Runs the AI Developer pipeline:
#     #     - Stage 1 (analysis_only=True): Perception + Reasoning (no commit, no doc, no learning)
#     #     - Stage 2 (full): Action (apply/commit), Documentation (TDD/README), Learning
#     #     """

#     #     print(f"\n🚀 [INIT] Running AI Developer pipeline for: {repo_url}")

#     #     # Ensure session memory is initialized
#     #     if not hasattr(self, "session_memory"):
#     #         self.session_memory = {}
#     #     if user_id not in self.session_memory:
#     #         self.session_memory[user_id] = {}

#     #     # ⚡ Reuse cached analysis for same repo if available
#     #     cached_entry = self.session_memory[user_id].get(repo_url)
#     #     if cached_entry and not commit_changes:
#     #         print(f"⚡ [CACHE] Using cached analysis for user '{user_id}'")
#     #         files, analyses = cached_entry
#     #         repo_name = self._parse_repo(repo_url)[1]
#     #     else:
#     #         # 🧠 Stage 1 — Perception
#     #         repo = self.fetch_repo(repo_url)
#     #         repo_name = repo["repo_name"]
#     #         files = repo["files"]

#     #         if not files:
#     #             return {
#     #                 "repo": repo_name,
#     #                 "analysis": [],
#     #                 "project_doc": "No code files found.",
#     #                 "user_learning": {},
#     #             }

#     #         # 🧩 Stage 2 — Reasoning (Analyze files)
#     #         analyses: List[Dict[str, Any]] = []
#     #         for f in files:
#     #             a = self.analyze_code(f["file_name"], f["content"])
#     #             a["file"] = f["file_name"]
#     #             analyses.append(a)

#     #         # Cache analysis for reuse
#     #         self.session_memory[user_id][repo_url] = (files, analyses)

#     #     # 🧾 If analysis_only mode — stop here (Stage 1 only)
#     #     if analysis_only:
#     #         print("🧠 [ANALYSIS-ONLY] Returning structured analysis (no commit, no doc, no learning).")
#     #         return {
#     #             "repo": repo_name,
#     #             "analysis": analyses
#     #         }

#     #     # ⚙️ Stage 3 — Action (Apply accepted suggestions & commit)
#     #     if commit_changes and accepted_suggestions:
#     #         print("\n[ACTION] 🧩 Applying developer-approved suggestions and committing changes...")
#     #         username, reponame = self._parse_repo(repo_url)
#     #         base_branch = self._get_default_branch(username, reponame)
#     #         feature_branch = f"codezen/fixes-{int(time.time())}"
#     #         self._ensure_feature_branch(username, reponame, base_branch, feature_branch)

#     #         grouped: Dict[str, List[Dict[str, Any]]] = {}

#     #         # ✅ Group accepted suggestions by file
#     #         for s in accepted_suggestions:
#     #             s_dict = s if isinstance(s, dict) else s.__dict__
#     #             accepted = s_dict.get("accepted")
#     #             file_name = s_dict.get("file_name")
#     #             comment = s_dict.get("comment", "")
#     #             if accepted and file_name:
#     #                 grouped.setdefault(file_name, []).append({
#     #                     "file_name": file_name,
#     #                     "comment": comment,
#     #                     "accepted": accepted
#     #                 })

#     #         # ✅ Apply AI suggestions and commit
#     #         for file_name, suggs in grouped.items():
#     #             if file_name.endswith(".ipynb"):
#     #                 print(f"⚠️ [NOTEBOOK] Skipping modification for {file_name} to prevent corruption.")
#     #                 continue

#     #             try:
#     #                 current_content, sha = self._get_file_content_and_sha(username, reponame, feature_branch, file_name)
#     #                 updated_content = self.apply_suggestions(file_name, current_content, suggs)

#     #                 if current_content.strip() == updated_content.strip():
#     #                     print(f"⚪ [AI-FIX] No changes detected for {file_name}. Skipping commit.")
#     #                     continue

#     #                 diff = difflib.unified_diff(
#     #                     current_content.splitlines(),
#     #                     updated_content.splitlines(),
#     #                     fromfile=f"a/{file_name}",
#     #                     tofile=f"b/{file_name}",
#     #                     lineterm=""
#     #                 )
#     #                 diff_text = "\n".join(diff)

#     #                 entry = next((a for a in analyses if a["file"] == file_name), None)
#     #                 if entry:
#     #                     entry["diff"] = diff_text
#     #                     entry["changes_applied"] = True

#     #                 message = f"AI-Fix: Applied {len(suggs)} developer-approved suggestion(s) to {file_name}"
#     #                 self._commit_file_update(username, reponame, feature_branch, file_name, updated_content, message, sha)

#     #             except Exception as e:
#     #                 print(f"❌ [ACTION] Failed to commit {file_name}: {e}")

#     #         print(f"✅ [ACTION] All changes committed to branch: {feature_branch}")
#     #         self._create_pull_request(username, reponame, feature_branch, base_branch)
#     #     else:
#     #         print("[ACTION] ⚪ Skipping commit — No accepted suggestions or commit disabled.")

#     #     # 🧠 Stage 4 — Learning (store accepted/rejected feedback)
#     #     self.update_user_preferences(user_id, accepted_suggestions or [], rejected_suggestions or [])

#     #     # 📘 Stage 5 — Documentation
#     #     print(f"[PROJECT DOCS] 🏗️ Generating {doc_format.upper()} documentation for {repo_name}")
#     #     project_doc = self.generate_project_docs(repo_name, files, analyses, format=doc_format)

#     #     print("✅ [COMPLETE] AI Developer pipeline executed successfully.")
#     #     return {
#     #         "repo": repo_name,
#     #         "analysis": analyses,
#     #         "project_doc": project_doc,
#     #         "user_learning": self.user_preferences.get(user_id, {}),
#     #     }
    
#     def run_full_pipeline(
#         self,
#         repo_url: str,
#         user_id: str = "default_user",
#         accepted_suggestions: Optional[List[Dict[str, Any]]] = None,
#         rejected_suggestions: Optional[List[Dict[str, Any]]] = None,
#         doc_format: str = "md",
#         commit_changes: bool = False,
#         analysis_only: bool = False,
#         use_cache: bool = True,
#         generate_docs: bool = True
#     ) -> Dict[str, Any]:
#         """
#         Clean 3-stage pipeline:
#          1️⃣ Perception + Reasoning (analysis)
#          2️⃣ Action (apply/commit accepted suggestions)
#          3️⃣ Documentation (TDD/README)
#         """

#         print(f"\n🚀 [INIT] Running AI Developer pipeline for: {repo_url}")

#         # Ensure session memory
#         if not hasattr(self, "session_memory"):
#             self.session_memory = {}
#         if user_id not in self.session_memory:
#             self.session_memory[user_id] = {}

#         # Normalize key for caching
#         # cache_key = repo_url.strip()
#         # cached_entry = self.session_memory[user_id].get(cache_key)
#         # normalized_user = "default_user"  # Always use the same cache for analysis reuse
#         # cache_key = repo_url.strip()
#         # # cached_entry = self.session_memory.get(normalized_user, {}).get(cache_key)
#         # cached_entry = self.session_memory.setdefault(normalized_user, {})[cache_key] = (files, analyses)
#         normalized_user = "default_user"  # one cache for all users
#         cache_key = repo_url.strip()
#         cached_entry = self.session_memory.get(normalized_user, {}).get(cache_key)


#         # 1️⃣ PERCEPTION + REASONING
#         if use_cache and cached_entry:
#             print(f"⚡ [CACHE] Using cached analysis for user '{user_id}'")
#             files, analyses = cached_entry
#             repo_name = self._parse_repo(repo_url)[1]
            
#             if commit_changes:
#                 print(f"⚡ [CACHE] Skipping reanalysis — using cached repo + analysis for commit stage.")

#         else:
#             repo = self.fetch_repo(repo_url)
#             repo_name = repo["repo_name"]
#             files = repo["files"]
#             if not files:
#                 return {
#                     "repo": repo_name,
#                     "analysis": [],
#                     "project_doc": "No code files found.",
#                     "user_learning": {},
#                 }

#             # Reasoning — analyze all code files
#             analyses: List[Dict[str, Any]] = []
#             for f in files:
#                 a = self.analyze_code(f["file_name"], f["content"])
#                 a["file"] = f["file_name"]
#                 analyses.append(a)

#             # Cache for reuse
#             # self.session_memory[user_id][cache_key] = (files, analyses)
#             # print(f"✅ [CACHE] Analysis cached for repo: {repo_name}")
#             self.session_memory.setdefault(normalized_user, {})[cache_key] = (files, analyses)
#             print(f"✅ [CACHE] Analysis cached for repo: {repo_name}")


#         # Early return if only analysis is needed
#         if analysis_only:
#             print("🧠 [ANALYSIS-ONLY] Returning structured analysis (no commit, no doc, no learning).")
#             return {"repo": repo_name, "analysis": analyses}

#         # 2️⃣ ACTION — Apply & Commit accepted suggestions
#         if commit_changes and accepted_suggestions:
#             print("\n[ACTION] 🧩 Applying developer-approved suggestions and committing changes...")
#             username, reponame = self._parse_repo(repo_url)
#             base_branch = self._get_default_branch(username, reponame)
#             feature_branch = f"codezen/fixes-{int(time.time())}"
#             self._ensure_feature_branch(username, reponame, base_branch, feature_branch)

#             grouped: Dict[str, List[Dict[str, Any]]] = {}
#             for s in accepted_suggestions:
#                 s_dict = s if isinstance(s, dict) else s.__dict__
#                 file_name = s_dict.get("file_name")
#                 comment = s_dict.get("comment", "")
#                 if s_dict.get("accepted") and file_name:
#                     grouped.setdefault(file_name, []).append({"comment": comment})

#             for file_name, suggs in grouped.items():
#                 if file_name.endswith(".ipynb"):
#                     print(f"⚠️ [NOTEBOOK] Skipping modification for {file_name} (to prevent corruption).")
#                     continue

#                 try:
#                     current_content, sha = self._get_file_content_and_sha(username, reponame, feature_branch, file_name)
#                     updated_content = self.apply_suggestions(file_name, current_content, suggs)

#                     if current_content.strip() == updated_content.strip():
#                         print(f"⚪ [AI-FIX] No changes detected for {file_name}. Skipping commit.")
#                         continue

#                     diff = "\n".join(
#                         difflib.unified_diff(
#                             current_content.splitlines(),
#                             updated_content.splitlines(),
#                             fromfile=f"a/{file_name}",
#                             tofile=f"b/{file_name}",
#                             lineterm=""
#                         )
#                     )

#                     # Store diff for frontend display
#                     entry = next((a for a in analyses if a["file"] == file_name), None)
#                     if entry:
#                         entry["diff"] = diff
#                         entry["changes_applied"] = True

#                     commit_message = f"AI-Fix: Applied {len(suggs)} suggestion(s) to {file_name}"
#                     self._commit_file_update(username, reponame, feature_branch, file_name, updated_content, commit_message, sha)
#                 except Exception as e:
#                     print(f"❌ [ACTION] Failed for {file_name}: {e}")

#             print(f"✅ [ACTION] All changes committed to branch: {feature_branch}")
#             self._create_pull_request(username, reponame, feature_branch, base_branch)
#         else:
#             print("[ACTION] ⚪ Skipping commit — No accepted suggestions or commit disabled.")

#         # 3️⃣ LEARNING
#         self.update_user_preferences(user_id, accepted_suggestions or [], rejected_suggestions or [])

#         # 4️⃣ DOCUMENTATION
#         if generate_docs:
#             print(f"[PROJECT DOCS] 🏗️ Generating {doc_format.upper()} documentation for {repo_name}")
#             project_doc = self.generate_project_docs(repo_name, files, analyses, format=doc_format)
#         else:
#             project_doc = None

#         print("✅ [COMPLETE] AI Developer pipeline executed successfully.")

#         # ✅ Clean final return structure
#         result = {
#             "repo": repo_name,
#             "project_doc": project_doc,
#             "user_learning": self.user_preferences.get(user_id, {}),
#             "stage": "apply" if commit_changes else "analysis"
#         }

#         if not commit_changes:
#             result["analysis"] = analyses
#         else:
#             # Only send diffs to frontend (applied file changes)
#             result["diffs"] = [a for a in analyses if a.get("changes_applied")]

#         return result


import json
import time
import requests
import os
import re
import base64
import difflib
import markdown
from fpdf import FPDF
from docx import Document
from tempfile import NamedTemporaryFile
import weasyprint
from typing import Dict, Any, List, Optional
from autogen import ConversableAgent
import google.generativeai as genai
from dotenv import load_dotenv

# Load environment variables (API keys, etc.)
load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# ----------------------------------------
# Global Configurations
# ----------------------------------------
COMMENT_PREFIX = {
    ".py": "# ", ".js": "// ", ".java": "// ", ".c": "// ", ".cpp": "// ",
    ".ts": "// ", ".html": "", ".css": " */"
}


class DeveloperAgent(ConversableAgent):
    """
    Unified AI Developer Agent for CodeZen. Handles GitHub interaction,
    Gemini-based code analysis, refactoring, and documentation generation.
    It extends Autogen's ConversableAgent.
    """

    def __init__(self, model_name: str = "gemini-2.5-flash"):
        # 🔵 STEP 1: INITIALIZATION AND TOKEN HANDLING
        super().__init__(name="DeveloperAgent", description="Unified AI Developer Agent for CodeZen.")
        self.model_name = model_name
        self.model = genai.GenerativeModel(model_name)
        self.user_preferences: Dict[str, Any] = {}
        self.github_token: Optional[str] = None  # ✅ Stores user token dynamically from OAuth
        # Cache for analysis results to prevent re-running LLM on the same repo
        self.session_memory: Dict[str, Dict[str, Any]] = {}
        print(f"✅ DeveloperAgent initialized with {model_name}")

    def set_dynamic_token(self, token: str):
        """Store GitHub token from frontend OAuth for current user session."""
        self.github_token = token
        print("🔐 Dynamic GitHub token set successfully.")

    # ----------------------------------------
    # GitHub Authentication Helpers
    # ----------------------------------------
    def _github_headers(self) -> Dict[str, str]:
        """Uses dynamic token from OAuth if available, otherwise falls back to env variable."""
        token = self.github_token or os.getenv("GITHUB_TOKEN")
        if not token:
            raise Exception("Missing GitHub token (neither dynamic nor environment).")
        return {"Authorization": f"token {token}", "Accept": "application/vnd.github+json"}

    def _validate_github_token(self) -> bool:
        """Validate current GitHub token; returns True if valid, False if expired/invalid."""
        try:
            headers = self._github_headers()
            response = requests.get("https://api.github.com/user", headers=headers)
            if response.status_code == 200:
                return True
            if response.status_code in [401, 403]:
                print("⚠️ [AUTH] GitHub token is invalid or expired. Clearing token.")
                self.github_token = None
                return False
            print(f"⚠️ [AUTH] Unexpected GitHub API status: {response.status_code}")
            return False
        except Exception as e:
            print(f"❌ [AUTH] Error validating token: {e}")
            return False

    def _parse_repo(self, repo_url: str) -> tuple[str, str]:
        """Extract username and reponame from a GitHub URL."""
        parts = repo_url.strip("/").split("/")
        # Assumes structure is typically user/repo
        if len(parts) < 2:
            raise ValueError("Invalid GitHub repository URL format.")
        return parts[-2], parts[-1].replace(".git", "")

    # ----------------------------------------
    # 1️⃣ PERCEPTION — Fetch repo recursively
    # ----------------------------------------
    def fetch_repo(self, repo_url: str) -> Dict[str, Any]:
        """Fetches the repository file structure and content."""
        print(f"\n[PERCEPTION] 🌐 Fetching repository structure from: {repo_url}")
        username, reponame = self._parse_repo(repo_url)

        headers = self._github_headers() if self.github_token or os.getenv("GITHUB_TOKEN") else {}

        # Get default branch
        meta_resp = requests.get(f"https://api.github.com/repos/{username}/{reponame}", headers=headers)
        meta_resp.raise_for_status()
        meta = meta_resp.json()
        default_branch = meta.get("default_branch", "main")
        print(f"[PERCEPTION] 🪶 Default branch: {default_branch}")

        # Get the full recursive tree
        tree_resp = requests.get(
            f"https://api.github.com/repos/{username}/{reponame}/git/trees/{default_branch}?recursive=1",
            headers=headers
        )
        tree_resp.raise_for_status()
        tree = tree_resp.json()

        files: List[Dict[str, str]] = []
        allowed = [".py", ".js", ".java", ".ts", ".cpp", ".c", ".html", ".css", ".ipynb"]

        for item in tree.get("tree", []):
            if item["type"] == "blob" and any(item["path"].endswith(ext) for ext in allowed):
                raw = f"https://raw.githubusercontent.com/{username}/{reponame}/{default_branch}/{item['path']}"

                content_resp = requests.get(raw)
                try:
                    content_resp.raise_for_status()
                    content = content_resp.text
                except requests.HTTPError as e:
                    print(f"⚠️ [FETCH] Failed to get content for {item['path']}: {e}")
                    continue

                # Special handling for Jupyter Notebooks
                if item["path"].endswith(".ipynb"):
                    try:
                        nb = json.loads(content)
                        cells = [
                            "\n".join(c.get("source", []))
                            for c in nb.get("cells", [])
                            if c.get("cell_type") == "code"
                        ]
                        content = "\n\n".join(cells)
                        print(f"📘 [NOTEBOOK] Extracted {len(cells)} code cells from {item['path']}")
                    except Exception as e:
                        print(f"⚠️ [NOTEBOOK] Failed to parse notebook {item['path']}: {e}")

                files.append({"file_name": item["path"], "content": content})

        print(f"✅ [PERCEPTION] Repo fetched — {len(files)} code files detected.")
        return {"repo_name": reponame, "files": files}

    # ----------------------------------------
    # 2️⃣ REASONING — Analyze each file with Gemini
    # ----------------------------------------
    def analyze_code(self, file_name: str, content: str) -> Dict[str, Any]:
        """Analyzes a single code file using the Gemini model."""
        print(f"\n[REASONING] 🧩 Analyzing file: {file_name}")
        prompt = f"""
        You are an expert AI code reviewer.
        Analyze the following code and return **ONLY valid JSON** with no extra text or markdown.

        The JSON object MUST follow this structure:
        {{
          "issues": ["list of problems or bugs (e.g., hardcoded secrets, security risks)"],
          "refactors": ["list of improvement suggestions (e.g., better structure, modularity)"],
          "summary": "short summary of what this file does"
        }}

        Code:
        ```
        {content}
        ```
        """
        try:
            # Generate content from the model
            out = self.model.generate_content(prompt).text.strip()

            try:
                # Attempt to parse directly
                parsed = json.loads(out)
            except json.JSONDecodeError:
                # Attempt to extract JSON from surrounding text/markdown
                match = re.search(r"\{.*\}", out, re.DOTALL)
                if match:
                    # Clean up common markdown/text wrappers
                    clean_json_str = match.group().replace('```json', '').replace('```', '').strip()
                    parsed = json.loads(clean_json_str)
                else:
                    # Fallback to a basic structure if no valid JSON is found
                    print("⚠️ [REASONING] Failed JSON parse. Falling back to summary string.")
                    parsed = {"issues": [], "refactors": [], "summary": out[:300]}

            print(f"✅ [REASONING] Parsed analysis for {file_name}")
            return parsed

        except Exception as e:
            print(f"❌ [REASONING] Error: {e}")
            return {"file": file_name, "issues": [], "refactors": [], "summary": f"Analysis failed: {str(e)}"}

    # ----------------------------------------
    # 3️⃣ ACTION — GitHub Commit Helpers
    # ----------------------------------------

    def _get_default_branch(self, username: str, reponame: str) -> str:
        """Gets the default branch name of a repository."""
        if not self._validate_github_token():
            raise Exception("GitHub token expired or invalid. Please re-login via frontend.")
        resp = requests.get(f"https://api.github.com/repos/{username}/{reponame}", headers=self._github_headers())
        resp.raise_for_status()
        return resp.json().get("default_branch", "main")

    def _ensure_feature_branch(self, username: str, reponame: str, base_branch: str, feature_branch: str):
        """Creates a new feature branch from the base branch if it doesn't exist."""
        if not self._validate_github_token():
            raise Exception("GitHub token expired or invalid. Please re-login via frontend.")

        # 1. Get SHA of the base branch
        ref_url = f"https://api.github.com/repos/{username}/{reponame}/git/ref/heads/{base_branch}"
        ref_resp = requests.get(ref_url, headers=self._github_headers())
        ref_resp.raise_for_status()
        base_sha = ref_resp.json()["object"]["sha"]

        # 2. Try to create the new feature branch
        create_resp = requests.post(
            f"https://api.github.com/repos/{username}/{reponame}/git/refs",
            headers=self._github_headers(),
            json={"ref": f"refs/heads/{feature_branch}", "sha": base_sha},
        )

        if create_resp.status_code == 201:
            print(f"✅ [ACTION] Created new branch: {feature_branch}")
        elif create_resp.status_code == 422 and "Reference already exists" in create_resp.text:
            print(f"⚠️ [ACTION] Branch {feature_branch} already exists (reusing).")
        else:
            print(f"❌ [ACTION] Failed to create branch: {create_resp.text}")
            create_resp.raise_for_status()

    def _get_file_content_and_sha(self, username: str, reponame: str, branch: str, path: str) -> tuple[str, str]:
        """Fetches file content and its current SHA for committing updates."""
        resp = requests.get(
            f"https://api.github.com/repos/{username}/{reponame}/contents/{path}?ref={branch}",
            headers=self._github_headers(),
        )
        resp.raise_for_status()
        data = resp.json()
        # GitHub content is Base64 encoded
        content = base64.b64decode(data["content"]).decode("utf-8", errors="ignore")
        return content, data["sha"]

    def _commit_file_update(self, username: str, reponame: str, branch: str, path: str, new_content: str, commit_message: str, sha: str):
        """Commits a file update to a specific branch."""
        if not self._validate_github_token():
            raise Exception("GitHub token expired or invalid. Please re-login via frontend.")

        encoded = base64.b64encode(new_content.encode("utf-8")).decode("utf-8")
        resp = requests.put(
            f"https://api.github.com/repos/{username}/{reponame}/contents/{path}",
            headers=self._github_headers(),
            json={"message": commit_message, "content": encoded, "branch": branch, "sha": sha},
        )

        if resp.status_code in [200, 201]:
            print(f"✅ [ACTION] Committed {path} → {branch}")
        elif resp.status_code == 401:
            print("❌ [AUTH] Invalid GitHub token during commit.")
            raise Exception("GitHub token expired or invalid. Please re-login.")
        else:
            resp.raise_for_status()

    def apply_suggestions(self, file_path: str, original_content: str, suggestions: List[Dict[str, Any]]) -> str:
        """Uses Gemini to apply accepted suggestions by rewriting the code."""
        print(f"\n[AI-FIX] ✏️ Applying {len(suggestions)} suggestion(s) to {file_path}")

        # Concatenate all developer-approved comments into one instruction block
        suggestions_text = "\n".join(
            [f"- {s.get('comment', '')}" for s in suggestions if s.get("comment")]
        )

        ext = os.path.splitext(file_path)[1]

        prompt = f"""
        You are an AI software engineer.
        Modify the following code based on the accepted refactor suggestions.
        Return **ONLY the updated code** (do not include explanations, comments, or markdown formatting like ```).

        File type: {ext}
        Accepted suggestions:
        {suggestions_text}

        Original code:
        {original_content}
        """

        try:
            result = self.model.generate_content(prompt)
            # Remove any surrounding markdown code block if present
            updated_code = result.text.strip().removeprefix(f"```{ext}").removeprefix("```").removesuffix("```").strip()

            if not updated_code:
                print(f"⚠️ [AI-FIX] No updated code returned for {file_path}. Keeping original.")
                return original_content

            print(f"✅ [AI-FIX] Successfully updated {file_path}")
            return updated_code
        except Exception as e:
            print(f"❌ [AI-FIX] Error applying suggestions to {file_path}: {e}")
            return original_content

    # ----------------------------------------
    # 4️⃣ DOCUMENTATION GENERATION (Enhanced)
    # ----------------------------------------
    @staticmethod
    def convert_doc(markdown_text: str, output_format: str = "pdf") -> str:
        """
        Convert Markdown text into either a PDF or DOCX file using external libraries.
        Returns the path to the generated file.
        """
        output_format = output_format.lower().strip()
        file = NamedTemporaryFile(delete=False, suffix=f".{output_format}")
        output_path = file.name

        if output_format == "pdf":
            html = markdown.markdown(markdown_text)
            try:
                # Preferred: Use weasyprint for better HTML/CSS handling
                weasyprint.HTML(string=html).write_pdf(output_path)
            except Exception:
                # Fallback: Use FPDF (less formatting, plain text layout)
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                # FPDF does not directly support Markdown; need to approximate
                for line in html.splitlines():
                    pdf.multi_cell(0, 10, line.strip())
                pdf.output(output_path)

        elif output_format == "docx":
            doc = Document()
            for line in markdown_text.split("\n"):
                doc.add_paragraph(line)
            doc.save(output_path)

        else:
            raise ValueError("Unsupported format. Choose 'pdf' or 'docx'.")

        return output_path

    def generate_project_docs(
        self,
        repo_name: str,
        files: List[Dict[str, str]],
        analyses: List[Dict[str, Any]],
        format: str = "pdf"
    ) -> str:
        """
        Generates a comprehensive Technical Design Document (TDD) using a
        structured prompt, and converts it to the desired format (Markdown, PDF, or DOCX).
        """
        print(f"\n[PROJECT DOCS] 🏗️ Generating {format.upper()} documentation for: {repo_name}")

        # Concatenate summaries for project context
        summarized = "\n".join(
            [f"- **{a.get('file', 'N/A')}** → {a.get('summary', '')}" for a in analyses if a.get("summary")]
        )

        project_prompt = f"""
        You are a senior software architect and documentation expert.
        Based on the following repository and analysis data, generate a **Technical Design Document (TDD)**
        for the project named **{repo_name}**.

        ---
        ### Document Characteristics
        - **Type:** Technical Design Document (TDD)
        - **Audience:** Technical stakeholders — AI/ML Team, Developers, Engineering Managers
        - **Purpose:** Explain system design, architecture, and engineering logic
        - **Tone:** Formal, structured, and explanatory
        - **Format:** Markdown with hierarchical numbering (1.0, 1.1, etc.)
        ---

        ### Standardized Section Template (Generate a detailed, multi-paragraph response for each main section)

        1. Project Overview
            1.1 Project Summary
            1.2 Problem Statement
            1.3 Business Use Case
            1.4 Target Users

        2. Functionality & Features
            2.1 Core Functionalities
            2.2 File or Data Flow Description

        3. Architecture & Workflow
            3.1 High-Level Architecture (Describe the system components and their relationships)
            3.2 Module Responsibilities (Briefly describe what each core file does)
            3.3 Execution Flow (The sequence of operations/calls)
        - Include a simple ASCII diagram representing the main flow. Example:
          ```
          [Client (Web/CLI)] → [API Gateway] → [main_service.py] → [data_storage.db]
          ```

        4. Technology Stack
            4.1 Programming Languages
            4.2 Libraries & Tools
            4.3 External Services or APIs

        5. Setup Instructions
            5.1 System Requirements (e.g., Python 3.9+, Node.js 18+)
            5.2 Installation & Configuration
            5.3 How to Run the Application

        6. Testing & Validation
            - Describe the testing strategy (e.g., unit tests, integration tests).
            - Include a small table of example test cases.

        7. Prompt Engineering & Strategy (Crucial for LLM-based projects)
            7.1 Prompt Design Principles (e.g., Chain-of-Thought, few-shot)
            7.2 Example Prompts
            7.3 Evaluation Logic

        8. Troubleshooting (Provide a few common issues and solutions)

        9. Future Scope (List at least three potential enhancements)

        10. Appendix
            10.1 Glossary (Define key technical terms)

        ---
        ### Repository Context (Base your answers on this data)
        {summarized}

        Generate the complete document in a clean Markdown format, formatted with clear numbering and subheadings,
        without any extraneous preamble or postamble text.
        """

        # Generate with Gemini
        try:
            res = self.model.generate_content(project_prompt)
            final_doc = res.text.strip()
            print("✅ [PROJECT DOCS] Raw documentation generated successfully.")

            # Conversion
            if format.lower() in ["pdf", "docx"]:
                print(f"📄 [CONVERT] Converting to {format.upper()} format...")
                output_path = self.convert_doc(final_doc, output_format=format)
                print(f"✅ [PROJECT DOCS] Saved as {output_path}")
                return output_path
            else:
                # Return raw Markdown with a wrapper for clarity
                return f"\n---\nPROJECT DOCUMENTATION: {repo_name}\n---\n\n{final_doc}\n\nGenerated by CodeZen AI Developer Agent"

        except Exception as e:
            print(f"❌ [PROJECT DOCS] Error: {e}")
            return f"Error generating documentation: {e}"


    # ----------------------------------------
    # 5️⃣ LEARNING — Store user feedback locally
    # ----------------------------------------
    def update_user_preferences(self, user_id: str, accepted: Optional[List[Dict[str, Any]]], rejected: Optional[List[Dict[str, Any]]]):
        """Stores accepted/rejected suggestions for future model fine-tuning or personalized analysis."""
        memory_file = ".codezen_memory.json"
        memory_data: Dict[str, Any] = {}

        # Load existing memory
        if os.path.exists(memory_file):
            try:
                with open(memory_file, "r", encoding="utf-8") as f:
                    memory_data = json.load(f)
            except Exception:
                memory_data = {} # Reset if file is corrupted

        if user_id not in memory_data:
            memory_data[user_id] = {"accepted": [], "rejected": []}

        current_time = time.time()

        # Process and structure the feedback data
        accepted_data = [
            {"file_name": s.get("file_name"), "comment": s.get("comment"), "timestamp": current_time}
            for s in accepted or [] if isinstance(s, dict)
        ]
        rejected_data = [
            {"file_name": s.get("file_name"), "comment": s.get("comment"), "timestamp": current_time}
            for s in rejected or [] if isinstance(s, dict)
        ]

        memory_data[user_id]["accepted"].extend(accepted_data)
        memory_data[user_id]["rejected"].extend(rejected_data)

        # Save updated memory
        try:
            with open(memory_file, "w", encoding="utf-8") as f:
                json.dump(memory_data, f, indent=2)
            self.user_preferences[user_id] = memory_data[user_id]
            print(f"🧠 [LEARNING] Updated user preferences for '{user_id}' — {len(accepted_data)} accepted, {len(rejected_data)} rejected suggestions.")
        except Exception as e:
            print(f"❌ [LEARNING] Failed to write memory file: {e}")

    # ----------------------------------------
    # 6️⃣ PULL REQUEST — Auto Create after AI commits
    # ----------------------------------------
    def _create_pull_request(self, username: str, reponame: str, feature_branch: str, base_branch: str) -> Optional[str]:
        """Automatically open a pull request on GitHub after AI commits."""
        try:
            pr_url = f"https://api.github.com/repos/{username}/{reponame}/pulls"
            headers = self._github_headers()
            pr_title = f"🤖 CodeZen AI Refactors — {feature_branch}"
            pr_body = (
                f"### Automated Code Improvements by CodeZen\n\n"
                f"CodeZen AI applied developer-approved refactors and improvements.\n\n"
                f"**Base Branch:** `{base_branch}`\n"
                f"**Feature Branch:** `{feature_branch}`\n\n"
                f"Please review the changes and merge if they look good. 🚀"
            )

            payload = {
                "title": pr_title,
                "head": feature_branch,
                "base": base_branch,
                "body": pr_body
            }

            response = requests.post(pr_url, headers=headers, json=payload)

            if response.status_code in [200, 201]:
                pr_data = response.json()
                print(f"✅ [PR] Pull Request Created: {pr_data.get('html_url')}")
                return pr_data.get("html_url")
            if response.status_code == 422:
                print("⚠️ [PR] A pull request for this branch already exists (or base/head error).")
            elif response.status_code == 401:
                print("❌ [PR] Unauthorized — GitHub token invalid or missing.")
            else:
                print(f"❌ [PR] Failed to create pull request: {response.text}")

        except Exception as e:
            print(f"❌ [PR] Error creating pull request: {e}")
        return None

    # ----------------------------------------
    # 7️⃣ FULL PIPELINE
    # ----------------------------------------
    def run_full_pipeline(
        self,
        repo_url: str,
        user_id: str = "default_user",
        accepted_suggestions: Optional[List[Dict[str, Any]]] = None,
        rejected_suggestions: Optional[List[Dict[str, Any]]] = None,
        doc_format: str = "md",
        commit_changes: bool = False,
        analysis_only: bool = False,
        use_cache: bool = True
    ) -> Dict[str, Any]:
        """
        Runs the full AI Developer pipeline (Perception, Reasoning, Action, Documentation, Learning).
        """

        print(f"\n🚀 [INIT] Running AI Developer pipeline for: {repo_url}")
        
        # Initialize session memory structure
        if user_id not in self.session_memory:
            self.session_memory[user_id] = {}
        
        repo_name = self._parse_repo(repo_url)[1] # Pre-parse repo name for error messages

        # Stage 1 & 2: Perception + Reasoning
        cached_entry = self.session_memory[user_id].get(repo_url)
        files: List[Dict[str, str]] = []
        analyses: List[Dict[str, Any]] = []

        if use_cache and cached_entry and not commit_changes:
            print(f"⚡ [CACHE] Using cached analysis for user '{user_id}'.")
            files, analyses = cached_entry
        else:
            try:
                repo = self.fetch_repo(repo_url)
                repo_name = repo["repo_name"]
                files = repo["files"]

                if not files:
                    return {
                        "repo": repo_name,
                        "analysis": [],
                        "project_doc": "No code files found.",
                        "pr_url": None,
                        "user_learning": {},
                    }

                # Reasoning (Analyze files)
                for f in files:
                    a = self.analyze_code(f["file_name"], f["content"])
                    a["file"] = f["file_name"]
                    analyses.append(a)

                # Cache analysis for reuse
                self.session_memory[user_id][repo_url] = (files, analyses)
            except Exception as e:
                print(f"❌ [PIPELINE] Initial fetching/analysis failed: {e}")
                return {
                    "repo": repo_name,
                    "analysis": [],
                    "project_doc": f"Error: {str(e)}",
                    "pr_url": None,
                    "user_learning": {},
                }


        if analysis_only:
            print("🧠 [ANALYSIS-ONLY] Returning structured analysis.")
            return {
                "repo": repo_name,
                "analysis": analyses,
                "pr_url": None,
            }
        
        # Initialize response variables
        pr_url: Optional[str] = None
        updated_files: List[str] = []
        
        # Stage 3: Action (Apply accepted suggestions & commit)
        if commit_changes and accepted_suggestions:
            print("\n[ACTION] 🧩 Applying developer-approved suggestions and committing changes...")
            username, reponame = self._parse_repo(repo_url)
            
            try:
                base_branch = self._get_default_branch(username, reponame)
                feature_branch = f"codezen/fixes-{int(time.time())}"
                self._ensure_feature_branch(username, reponame, base_branch, feature_branch)
            except Exception as e:
                print(f"❌ [ACTION] Branching failed: {e}")
                commit_changes = False # Disable commit steps

            if commit_changes:
                # Group accepted suggestions by file
                grouped: Dict[str, List[Dict[str, Any]]] = {}
                for s in accepted_suggestions:
                    file_name = s.get("file_name")
                    if file_name:
                        if file_name not in grouped:
                            grouped[file_name] = []
                        grouped[file_name].append(s)
                
                # Iterate through files that need updates
                for file_name, suggestions in grouped.items():
                    try:
                        # 1. Get current content and SHA
                        original_content, sha = self._get_file_content_and_sha(
                            username, reponame, base_branch, file_name
                        )

                        # 2. Apply suggestions via LLM
                        updated_content = self.apply_suggestions(
                            file_name, original_content, suggestions
                        )
                        
                        # 3. Commit the change if content actually changed
                        if original_content.strip() != updated_content.strip():
                            commit_msg = f"feat(codezen): apply {len(suggestions)} refactors to {file_name}"
                            self._commit_file_update(
                                username, reponame, feature_branch, file_name,
                                updated_content, commit_msg, sha
                            )
                            updated_files.append(file_name)
                        else:
                            print(f"⚠️ [ACTION] No change detected after AI fix for {file_name}. Skipping commit.")
                            
                    except Exception as e:
                        print(f"❌ [ACTION] Failed to process/commit {file_name}: {e}")
                
                # 4. Create Pull Request if commits were made
                if updated_files:
                    print(f"🎉 [ACTION] Commits made for {len(updated_files)} files.")
                    pr_url = self._create_pull_request(username, reponame, feature_branch, base_branch)
                else:
                    print("⚠️ [ACTION] No files were modified; skipping Pull Request creation.")


        # Stage 4: Documentation (Generate TDD/README)
        project_doc = self.generate_project_docs(
            repo_name, files, analyses, format=doc_format
        )
        
        # Stage 5: Learning (Store feedback)
        self.update_user_preferences(user_id, accepted_suggestions, rejected_suggestions)
        
        # Final result structure
        return {
            "repo": repo_name,
            "analysis": analyses,
            "project_doc": project_doc,
            "pr_url": pr_url,
            "user_learning": self.user_preferences.get(user_id, {})
        }

    # Alias for backward compatibility (if needed)
    apply_changes_and_generate_docs = run_full_pipeline
