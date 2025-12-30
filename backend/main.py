# from fastapi import FastAPI, Query, HTTPException
# from fastapi.middleware.cors import CORSMiddleware
# from pydantic import BaseModel
# import requests
# import os
# import google.generativeai as genai
# from dotenv import load_dotenv
# from documentation_agent import DocumentationAgent
# import time
# from fastapi.responses import FileResponse
# from reportlab.lib.pagesizes import A4
# from reportlab.pdfgen import canvas
# import tempfile
# from typing import Optional
# # -------------------------
# # Load environment variables
# # -------------------------
# print("[BOOT] Loading environment variables...")
# load_dotenv()
# GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

# if not GEMINI_API_KEY:
#     raise Exception("⚠️ GEMINI_API_KEY not found in .env file")

# genai.configure(api_key=GEMINI_API_KEY)
# print("✅ GEMINI_API_KEY configured successfully.")

# # -------------------------
# # Initialize FastAPI app
# # -------------------------
# app = FastAPI(title="CodeZen Backend", version="1.0")

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# print("🚀 FastAPI app initialized successfully.")

# # -------------------------
# # Root endpoint
# # -------------------------
# @app.get("/")
# def home():
#     return {"message": "CodeZen Backend is running 🚀"}

# # -------------------------
# # Fetch repository files
# # -------------------------
# @app.get("/fetch-repo")
# def fetch_repo(repo_url: str = Query(..., description="GitHub repository URL")):
#     print(f"\n[FETCH] Fetching repo from: {repo_url}")
#     start_time = time.time()

#     try:
#         parts = repo_url.strip("/").split("/")
#         username, reponame = parts[-2], parts[-1]
#         api_url = f"https://api.github.com/repos/{username}/{reponame}/contents"

#         print(f"[DEBUG] GitHub API URL: {api_url}")
#         response = requests.get(api_url)

#         if response.status_code != 200:
#             print(f"❌ Failed to fetch repo contents. Status: {response.status_code}")
#             return {"error": "Failed to fetch repo contents"}

#         repo_data = response.json()
#         code_files = []

#         for file in repo_data:
#             if file["type"] == "file" and file["name"].endswith(".py"):
#                 print(f"[INFO] Found Python file: {file['name']}")
#                 file_content = requests.get(file["download_url"]).text
#                 code_files.append({
#                     "file_name": file["name"],
#                     "content": file_content
#                 })

#         elapsed = round(time.time() - start_time, 2)
#         print(f"✅ Repo fetched successfully: {reponame} (Files: {len(code_files)}, Time: {elapsed}s)")

#         return {
#             "repo_name": reponame,
#             "total_files": len(code_files),
#             "files": code_files
#         }

#     except Exception as e:
#         print(f"❌ Exception while fetching repo: {e}")
#         return {"error": str(e)}

# # -------------------------
# # Initialize Agent
# # -------------------------
# print("[BOOT] Initializing Documentation Agent...")
# agent = DocumentationAgent()
# print("✅ Documentation Agent ready for requests.\n")

# # ---------- Models ----------
# class File(BaseModel):
#     file_name: str
#     content: str

# class FilesRequest(BaseModel):
#     files: list[File]

# # -------------------------
# # Generate Docs Endpoint
# # -------------------------
# @app.post("/generate-docs")
# async def generate_docs(request: FilesRequest):
#     print(f"\n[REQUEST] Received request to generate docs for {len(request.files)} files.")
#     print(f"[TRACE] Starting doc generation process for {len(request.files)} file(s)...")
#     start_time = time.time()

#     try:
#         docs_output = []
#         for file in request.files:
#             print(f"[PROCESS] Handling file: {file.file_name}")
#             doc_text = agent.generate_docs_for_file(file.file_name, file.content)
#             docs_output.append({
#                 "file_name": file.file_name,
#                 "documentation": doc_text
#             })

#         elapsed = round(time.time() - start_time, 2)
#         print(f"✅ All documentation generated successfully in {elapsed}s\n")
#         return {"documentation": docs_output}

#     except Exception as e:
#         print(f"❌ Error in /generate-docs endpoint: {e}")
#         raise HTTPException(status_code=500, detail=str(e))


# # -------------------------
# # Export Documentation Endpoint
# # -------------------------
# @app.post("/export-docs")
# async def export_docs(format: str, request: Optional[FilesRequest] = None):
#     """
#     Export generated documentation in desired format:
#     pdf | md | txt | docx | html
#     """
#     print(f"\n[EXPORT] Exporting documentation in format: {format.upper()}")
#     print(f"[TRACE] Starting export in {format.upper()} format...")

#     # ✅ Validate request
#     if not request or not request.files:
#         raise HTTPException(status_code=400, detail="No files provided for export")

#     try:
#         docs_output = []
#         for file in request.files:
#             print(f"[TRACE] Now exporting: {file.file_name}")
#             doc_text = agent.generate_docs_for_file(file.file_name, file.content)
#             docs_output.append({
#                 "file_name": file.file_name,
#                 "documentation": doc_text
#             })

#         # -------------------------
#         # PDF EXPORT
#         # -------------------------
#         if format.lower() == "pdf":
#             from reportlab.pdfgen import canvas
#             from reportlab.lib.pagesizes import A4
#             pdf_path = tempfile.mktemp(suffix=".pdf")
#             c = canvas.Canvas(pdf_path, pagesize=A4)
#             width, height = A4
#             y = height - 50

#             for doc in docs_output:
#                 c.setFont("Helvetica-Bold", 14)
#                 c.drawString(50, y, f"📄 {doc['file_name']}")
#                 y -= 20
#                 c.setFont("Helvetica", 11)

#                 for line in doc["documentation"].splitlines():
#                     if y < 50:
#                         c.showPage()
#                         y = height - 50
#                         c.setFont("Helvetica", 11)
#                     c.drawString(50, y, line[:110])
#                     y -= 15
#                 y -= 30

#             c.save()
#             print(f"✅ PDF generated successfully: {pdf_path}")
#             return FileResponse(pdf_path, media_type="application/pdf", filename="CodeZen_Docs.pdf")

#         # -------------------------
#         # MARKDOWN EXPORT
#         # -------------------------
#         elif format.lower() == "md":
#             md_path = tempfile.mktemp(suffix=".md")
#             with open(md_path, "w", encoding="utf-8") as f:
#                 for doc in docs_output:
#                     f.write(f"# {doc['file_name']}\n\n{doc['documentation']}\n\n---\n\n")
#             print(f"✅ Markdown file generated: {md_path}")
#             return FileResponse(md_path, media_type="text/markdown", filename="CodeZen_Docs.md")

#         # -------------------------
#         # TEXT EXPORT
#         # -------------------------
#         elif format.lower() == "txt":
#             txt_path = tempfile.mktemp(suffix=".txt")
#             with open(txt_path, "w", encoding="utf-8") as f:
#                 for doc in docs_output:
#                     f.write(f"{doc['file_name']}\n{'='*40}\n{doc['documentation']}\n\n")
#             print(f"✅ Text file generated: {txt_path}")
#             return FileResponse(txt_path, media_type="text/plain", filename="CodeZen_Docs.txt")

#         # -------------------------
#         # DOCX EXPORT
#         # -------------------------
#         elif format.lower() == "docx":
#             from docx import Document
#             docx_path = tempfile.mktemp(suffix=".docx")
#             document = Document()
#             document.add_heading("CodeZen Documentation", level=1)

#             for doc in docs_output:
#                 document.add_heading(doc["file_name"], level=2)
#                 document.add_paragraph(doc["documentation"])
#                 document.add_page_break()

#             document.save(docx_path)
#             print(f"✅ Word (DOCX) file generated: {docx_path}")
#             return FileResponse(
#                 docx_path,
#                 media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#                 filename="CodeZen_Docs.docx"
#             )

#         # -------------------------
#         # HTML EXPORT
#         # -------------------------
#         elif format.lower() == "html":
#             html_path = tempfile.mktemp(suffix=".html")
#             with open(html_path, "w", encoding="utf-8") as f:
#                 f.write("<html><head><title>CodeZen Documentation</title></head><body>")
#                 f.write("<h1>CodeZen Documentation</h1>")
#                 for doc in docs_output:
#                     f.write(f"<h2>{doc['file_name']}</h2>")
#                     formatted = doc["documentation"].replace("\n", "<br>")
#                     f.write(f"<p>{formatted}</p><hr>")
#                 f.write("</body></html>")
#             print(f"✅ HTML file generated: {html_path}")
#             return FileResponse(html_path, media_type="text/html", filename="CodeZen_Docs.html")

#         # -------------------------
#         # INVALID FORMAT
#         # -------------------------
#         else:
#             print("⚠️ Invalid format received.")
#             raise HTTPException(status_code=400, detail="Invalid format. Use 'pdf', 'md', 'txt', 'docx', or 'html'.")

#     except Exception as e:
#         print(f"❌ Error exporting docs: {e}")
#         raise HTTPException(status_code=500, detail=str(e))


# from fastapi import FastAPI, Query, HTTPException
# from fastapi.middleware.cors import CORSMiddleware
# from pydantic import BaseModel
# import requests, os, json, time, tempfile
# import google.generativeai as genai
# from fastapi.responses import FileResponse
# from dotenv import load_dotenv
# from typing import Optional
# from documentation_agent import DocumentationAgent
# from developer_agent import AIDeveloperAgent # 👈 new agent
# from fastapi import Body

# # -------------------------
# # Load environment variables
# # -------------------------
# print("[BOOT] Loading environment variables...")
# load_dotenv()
# GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

# if not GEMINI_API_KEY:
#     raise Exception("⚠️ GEMINI_API_KEY not found in .env file")

# genai.configure(api_key=GEMINI_API_KEY)
# print("✅ GEMINI_API_KEY configured successfully.")

# # -------------------------
# # Initialize FastAPI app
# # -------------------------
# app = FastAPI(title="CodeZen Backend", version="2.0")

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# print("🚀 FastAPI app initialized successfully.")

# # -------------------------
# # Root Endpoint
# # -------------------------
# @app.get("/")
# def home():
#     return {"message": "CodeZen Backend is running 🚀"}

# # -------------------------
# # Fetch Repository Files (Multi-language)
# # -------------------------
# @app.get("/fetch-repo")
# def fetch_repo(repo_url: str = Query(..., description="GitHub repository URL")):
#     print(f"\n[FETCH] Fetching repo from: {repo_url}")
#     start_time = time.time()

#     try:
#         parts = repo_url.strip("/").split("/")
#         username, reponame = parts[-2], parts[-1]
#         api_url = f"https://api.github.com/repos/{username}/{reponame}/contents"

#         response = requests.get(api_url)
#         if response.status_code != 200:
#             raise HTTPException(status_code=400, detail="Failed to fetch repo contents")

#         repo_data = response.json()
#         code_files = []

#         allowed_ext = [".py", ".js", ".java", ".ts", ".cpp", ".c", ".html", ".css", ".json", ".xml", ".md"]

#         for file in repo_data:
#             if file["type"] == "file" and any(file["name"].endswith(ext) for ext in allowed_ext):
#                 file_content = requests.get(file["download_url"]).text
#                 code_files.append({
#                     "file_name": file["name"],
#                     "content": file_content
#                 })

#         elapsed = round(time.time() - start_time, 2)
#         print(f"✅ Repo fetched successfully ({len(code_files)} files, {elapsed}s)")
#         return {
#             "repo_name": reponame,
#             "total_files": len(code_files),
#             "files": code_files
#         }

#     except Exception as e:
#         print(f"❌ Error fetching repo: {e}")
#         raise HTTPException(status_code=500, detail=str(e))

# # -------------------------
# # Initialize Agents
# # -------------------------
# print("[BOOT] Initializing Agents...")
# agent = DocumentationAgent()
# developer_agent = AIDeveloperAgent()
# print("✅ Agents ready.\n")

# # ---------- Models ----------
# class File(BaseModel):
#     file_name: str
#     content: str

# class FilesRequest(BaseModel):
#     files: list[File]

# class SuggestionResponse(BaseModel):
#     file_name: str
#     updated_content: Optional[str] = None
#     accepted: bool
#     comment: str

# # -------------------------
# # Analyze Code Endpoint (Perception + Reasoning)
# # -------------------------
# @app.post("/analyze-code")
# async def analyze_code(request: FilesRequest):
#     print(f"\n[AI] Starting code analysis for {len(request.files)} file(s)...")
#     try:
#         all_results = []
#         for file in request.files:
#             result = developer_agent.analyze_code(file.file_name, file.content)
#             all_results.append(result)

#         print("✅ Code analysis completed successfully.")
#         return {"analysis_results": all_results}
#     except Exception as e:
#         print(f"❌ Error in /analyze-code: {e}")
#         raise HTTPException(status_code=500, detail=str(e))

# # -------------------------
# # Apply Suggestions Endpoint (Action Stage)
# # -------------------------
# @app.post("/apply-suggestions")
# async def apply_suggestions(repo_url: str, suggestions: list[SuggestionResponse]):
#     print("\n[ACTION] Processing developer responses...")
#     try:
#         accepted_changes = [s for s in suggestions if s.accepted]

#         if not accepted_changes:
#             print("ℹ️ No accepted changes. Skipping commit.")
#         else:
#             print(f"📝 {len(accepted_changes)} changes accepted. Committing updates...")
#             # (Pseudo-commit logic)
#             for s in accepted_changes:
#                 print(f"→ Updating {s.file_name} | Comment: {s.comment}")

#         print("✅ Suggestions processed successfully.")
#         return {"message": "Suggestions processed successfully", "accepted": len(accepted_changes)}

#     except Exception as e:
#         print(f"❌ Error in /apply-suggestions: {e}")
#         raise HTTPException(status_code=500, detail=str(e))

# # -------------------------
# # Generate Documentation Endpoint
# # -------------------------
# @app.post("/generate-docs")
# async def generate_docs(request: FilesRequest):
#     print(f"\n[DOCS] Generating documentation for {len(request.files)} files...")
#     start_time = time.time()

#     try:
#         docs_output = []
#         for file in request.files:
#             doc_text = agent.generate_docs_for_file(file.file_name, file.content)
#             docs_output.append({
#                 "file_name": file.file_name,
#                 "documentation": doc_text
#             })

#         elapsed = round(time.time() - start_time, 2)
#         print(f"✅ Docs generated successfully in {elapsed}s.")
#         return {"documentation": docs_output}
#     except Exception as e:
#         print(f"❌ Error in /generate-docs: {e}")
#         raise HTTPException(status_code=500, detail=str(e))

# # -------------------------
# # Export Documentation Endpoint
# # -------------------------
# @app.post("/export-docs")
# async def export_docs(format: str, request: Optional[FilesRequest] = None):
#     if not request or not request.files:
#         raise HTTPException(status_code=400, detail="No files provided for export")

#     print(f"\n[EXPORT] Exporting documentation in format: {format.upper()}")

#     docs_output = []
#     for file in request.files:
#         doc_text = agent.generate_docs_for_file(file.file_name, file.content)
#         docs_output.append({
#             "file_name": file.file_name,
#             "documentation": doc_text
#         })

#     # ---------- Export Formats ----------
#     if format.lower() == "pdf":
#         from reportlab.pdfgen import canvas
#         from reportlab.lib.pagesizes import A4
#         pdf_path = tempfile.mktemp(suffix=".pdf")
#         c = canvas.Canvas(pdf_path, pagesize=A4)
#         width, height = A4
#         y = height - 50

#         for doc in docs_output:
#             c.setFont("Helvetica-Bold", 14)
#             c.drawString(50, y, f"{doc['file_name']}")
#             y -= 20
#             c.setFont("Helvetica", 11)

#             for line in doc["documentation"].splitlines():
#                 if y < 50:
#                     c.showPage()
#                     y = height - 50
#                 c.drawString(50, y, line[:110])
#                 y -= 15
#             y -= 30

#         c.save()
#         return FileResponse(pdf_path, media_type="application/pdf", filename="CodeZen_Docs.pdf")

#     elif format.lower() == "md":
#         md_path = tempfile.mktemp(suffix=".md")
#         with open(md_path, "w", encoding="utf-8") as f:
#             for doc in docs_output:
#                 f.write(f"# {doc['file_name']}\n\n{doc['documentation']}\n\n---\n")
#         return FileResponse(md_path, media_type="text/markdown", filename="CodeZen_Docs.md")

#     elif format.lower() == "docx":
#         from docx import Document
#         docx_path = tempfile.mktemp(suffix=".docx")
#         document = Document()
#         document.add_heading("CodeZen Documentation", level=1)
#         for doc in docs_output:
#             document.add_heading(doc["file_name"], level=2)
#             document.add_paragraph(doc["documentation"])
#         document.save(docx_path)
#         return FileResponse(docx_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="CodeZen_Docs.docx")

#     elif format.lower() == "html":
#         html_path = tempfile.mktemp(suffix=".html")
#         with open(html_path, "w", encoding="utf-8") as f:
#             f.write("<html><body><h1>CodeZen Documentation</h1>")
#             for doc in docs_output:
#                 safe_html = doc['documentation'].replace('\n','br')
#                 f.write(f"<h2>{doc['file_name']}</h2><p>{safe_html}</p><hr>")
#             f.write("</body></html>")
#         return FileResponse(html_path, media_type="text/html", filename="CodeZen_Docs.html")

#     else:
#         raise HTTPException(status_code=400, detail="Invalid format. Use pdf, md, docx, html.")
    
# @app.post("/run-agent")
# async def run_agent(
#     repo_url: str = Body(..., embed=True),
#     accepted_suggestions: Optional[list[SuggestionResponse]] = Body(default=None)
# ):
#     """
#     Run full CodeZen pipeline for given repo.
#     Fetch → Analyze → Apply suggestions → Generate Docs.
#     """
#     print(f"\n🚀 [RUN] Starting full agent pipeline for repo: {repo_url}")
#     try:
#         result = developer_agent.run_full_pipeline(repo_url, accepted_suggestions)
#         print("✅ Full pipeline executed successfully.")
#         return result
#     except Exception as e:
#         print(f"❌ Error in /run-agent: {e}")
#         raise HTTPException(status_code=500, detail=str(e))

# # main.py
# from fastapi import FastAPI, HTTPException, Body
# from fastapi.middleware.cors import CORSMiddleware
# from developer_agent import DeveloperAgent
# from pydantic import BaseModel
# from typing import Optional

# app = FastAPI(title="CodeZen Backend", version="3.0")

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# print("🚀 FastAPI Initialized")

# # Initialize the unified DeveloperAgent
# developer_agent = DeveloperAgent()

# class SuggestionResponse(BaseModel):
#     file_name: str
#     updated_content: Optional[str] = None
#     accepted: bool
#     comment: Optional[str] = None


# @app.post("/run-agent")
# async def run_agent(
#     repo_url: str = Body(..., embed=True),
#     accepted_suggestions: Optional[list[SuggestionResponse]] = Body(default=None)
# ):
#     """Run the full CodeZen pipeline via Autogen DeveloperAgent."""
#     try:
#         result = developer_agent.run_full_pipeline(repo_url, accepted_suggestions)
#         return {"status": "success", "result": result}
#     except Exception as e:
#         print(f"❌ Error running agent: {e}")
#         raise HTTPException(status_code=500, detail=str(e))

# # main.py
# from fastapi import FastAPI, HTTPException, Body
# from fastapi.middleware.cors import CORSMiddleware
# from developer_agent import DeveloperAgent
# from pydantic import BaseModel
# from typing import Optional, List
# from fastapi.responses import FileResponse
# import os, tempfile

# app = FastAPI(title="CodeZen Backend", version="3.1")

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# print("🚀 FastAPI Initialized")

# # Initialize the unified DeveloperAgent
# developer_agent = DeveloperAgent()


# # ----------------------------------------
# # Models
# # ----------------------------------------
# class SuggestionResponse(BaseModel):
#     file_name: str
#     updated_content: Optional[str] = None
#     accepted: bool
#     comment: Optional[str] = None


# class FileDoc(BaseModel):
#     file_name: str
#     documentation: str


# class ExportRequest(BaseModel):
#     format: str
#     docs: List[FileDoc]


# # ----------------------------------------
# # Run Full AI Agent Pipeline
# # ----------------------------------------
# @app.post("/run-agent")
# async def run_agent(
#     repo_url: str = Body(..., embed=True),
#     accepted_suggestions: Optional[List[SuggestionResponse]] = Body(default=None),
#     rejected_suggestions: Optional[List[SuggestionResponse]] = Body(default=None),
#     doc_format: str = Body("md", embed=True),
#     commit_changes: bool = Body(False, embed=True),
#     user_id: str = Body("default_user", embed=True)
# ):
#     """
#     Run the full CodeZen pipeline:
#     - Fetch repo
#     - Analyze code
#     - Apply suggestions (if any)
#     - Commit (optional)
#     - Generate documentation
#     """
#     try:
#         result = developer_agent.run_full_pipeline(
#             repo_url=repo_url,
#             user_id=user_id,
#             accepted_suggestions=accepted_suggestions,
#             rejected_suggestions=rejected_suggestions,
#             doc_format=doc_format,
#             commit_changes=commit_changes,
#         )
#         return {"status": "success", "result": result}
#     except Exception as e:
#         print(f"❌ Error running agent: {e}")
#         raise HTTPException(status_code=500, detail=str(e))


# # ----------------------------------------
# # Export Documentation (Download or Live Preview)
# # ----------------------------------------
# @app.post("/export-docs")
# async def export_docs(request: ExportRequest, preview: bool = False):
#     """
#     Export documentation in desired format (pdf | md | txt | docx | html).
#     If preview=True and format='html', returns HTML content for live preview.
#     """
#     try:
#         print(f"\n[EXPORT] Exporting docs in format: {request.format.upper()}")
#         docs_output = [{"file_name": d.file_name, "documentation": d.documentation} for d in request.docs]

#         format = request.format.lower()
#         path = None

#         # PDF Export
#         if format == "pdf":
#             from reportlab.pdfgen import canvas
#             from reportlab.lib.pagesizes import A4
#             path = tempfile.mktemp(suffix=".pdf")
#             c = canvas.Canvas(path, pagesize=A4)
#             width, height = A4
#             y = height - 50
#             for doc in docs_output:
#                 c.setFont("Helvetica-Bold", 14)
#                 c.drawString(50, y, f"{doc['file_name']}")
#                 y -= 20
#                 c.setFont("Helvetica", 11)
#                 for line in doc["documentation"].splitlines():
#                     if y < 50:
#                         c.showPage()
#                         y = height - 50
#                         c.setFont("Helvetica", 11)
#                     c.drawString(50, y, line[:110])
#                     y -= 15
#                 y -= 30
#             c.save()

#         # Markdown Export
#         elif format == "md":
#             path = tempfile.mktemp(suffix=".md")
#             with open(path, "w", encoding="utf-8") as f:
#                 for doc in docs_output:
#                     f.write(f"# {doc['file_name']}\n\n{doc['documentation']}\n\n---\n")

#         # Text Export
#         elif format == "txt":
#             path = tempfile.mktemp(suffix=".txt")
#             with open(path, "w", encoding="utf-8") as f:
#                 for doc in docs_output:
#                     f.write(f"{doc['file_name']}\n{'='*40}\n{doc['documentation']}\n\n")

#         # DOCX Export
#         elif format == "docx":
#             from docx import Document
#             path = tempfile.mktemp(suffix=".docx")
#             document = Document()
#             document.add_heading("CodeZen Documentation", level=1)
#             for doc in docs_output:
#                 document.add_heading(doc["file_name"], level=2)
#                 document.add_paragraph(doc["documentation"])
#                 document.add_page_break()
#             document.save(path)

#         # HTML Export / Live Preview
#         elif format == "html":
#             path = tempfile.mktemp(suffix=".html")
#             html_content = (
#                 "<html><head><title>CodeZen Documentation</title></head><body>"
#                 "<h1>CodeZen Documentation</h1>"
#             )
#             for doc in docs_output:
#                 formatted = doc["documentation"].replace("\n", "<br>")
#                 html_content += f"<h2>{doc['file_name']}</h2><p>{formatted}</p><hr>"
#             html_content += "</body></html>"

#             with open(path, "w", encoding="utf-8") as f:
#                 f.write(html_content)

#             if preview:
#                 return {"status": "success", "html": html_content}

#         else:
#             raise HTTPException(status_code=400, detail="Invalid format. Use pdf, md, txt, docx, or html.")

#         filename = f"CodeZen_Docs.{format}"
#         media_map = {
#             "pdf": "application/pdf",
#             "md": "text/markdown",
#             "txt": "text/plain",
#             "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#             "html": "text/html",
#         }

#         print(f"✅ Export successful → {path}")
#         return FileResponse(path, media_type=media_map.get(format, "application/octet-stream"), filename=filename)

#     except Exception as e:
#         print(f"❌ Error exporting docs: {e}")
#         raise HTTPException(status_code=500, detail=str(e))


# from fastapi import FastAPI, HTTPException, Body
# from fastapi.middleware.cors import CORSMiddleware
# from fastapi.responses import FileResponse, RedirectResponse, JSONResponse
# from developer_agent import DeveloperAgent
# from pydantic import BaseModel
# from typing import Optional, List
# import os, tempfile, requests

# app = FastAPI(title="CodeZen Backend", version="3.5")

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],  # ⚠️ For production, limit to your frontend URL
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# print("🚀 FastAPI Initialized")
# developer_agent = DeveloperAgent()

# # ========== MODELS ==========
# class SuggestionResponse(BaseModel):
#     file_name: str
#     updated_content: Optional[str] = None
#     accepted: bool
#     comment: Optional[str] = None

# class FileDoc(BaseModel):
#     file_name: str
#     documentation: str

# class ExportRequest(BaseModel):
#     format: str
#     docs: List[FileDoc]


# # ========== GITHUB OAUTH ==========
# GITHUB_CLIENT_ID = os.getenv("GITHUB_CLIENT_ID")
# GITHUB_CLIENT_SECRET = os.getenv("GITHUB_CLIENT_SECRET")
# REDIRECT_URI = "http://localhost:8000/oauth/callback"
# FRONTEND_URL = "http://localhost:3000"

# @app.get("/login/github")
# def github_login():
#     """
#     Step 1: Redirect user to GitHub OAuth consent screen
#     """
#     github_auth_url = (
#         "https://github.com/login/oauth/authorize"
#         f"?client_id={GITHUB_CLIENT_ID}"
#         f"&redirect_uri={REDIRECT_URI}"
#         "&scope=repo,user:email"
#         "&allow_signup=true"
#     )
#     print("🔗 Redirecting user to GitHub OAuth:", github_auth_url)
#     return RedirectResponse(url=github_auth_url)


# @app.get("/oauth/callback")
# def github_callback(code: str):
#     """
#     Step 2: GitHub redirects here with temporary 'code' — exchange for access_token.
#     """
#     try:
#         token_url = "https://github.com/login/oauth/access_token"
#         payload = {
#             "client_id": GITHUB_CLIENT_ID,
#             "client_secret": GITHUB_CLIENT_SECRET,
#             "code": code,
#         }
#         headers = {"Accept": "application/json"}
#         response = requests.post(token_url, json=payload, headers=headers)
#         token_data = response.json()

#         if "access_token" not in token_data:
#             raise Exception(token_data.get("error_description", "Failed to get token"))

#         access_token = token_data["access_token"]
#         print(f"✅ GitHub OAuth success — Token acquired for user: {access_token[:8]}...")

#         # Redirect back to frontend with token
#         return RedirectResponse(f"{FRONTEND_URL}?token={access_token}")

#     except Exception as e:
#         print(f"❌ OAuth callback error: {e}")
#         return JSONResponse({"error": str(e)}, status_code=500)


# # ========== STAGE 1+2: ANALYSIS ==========
# @app.post("/run-analysis")
# async def run_analysis(repo_url: str = Body(..., embed=True)):
#     try:
#         result = developer_agent.run_analysis_only(repo_url=repo_url)
#         return {"status": "success", "result": result}
#     except Exception as e:
#         print(f"❌ Error in run-analysis: {e}")
#         raise HTTPException(status_code=500, detail=str(e))


# # ========== STAGE 3+4: APPLY CHANGES + DOCS ==========
# @app.post("/apply-changes")
# async def apply_changes(
#     repo_url: str = Body(..., embed=True),
#     accepted_suggestions: Optional[List[SuggestionResponse]] = Body(default=None),
#     rejected_suggestions: Optional[List[SuggestionResponse]] = Body(default=None),
#     doc_format: str = Body("md", embed=True),
#     commit_changes: bool = Body(False, embed=True),
#     user_id: str = Body("default_user", embed=True),
#     auth_token: Optional[str] = Body(None, embed=True),
# ):
#     """
#     Stage 3+4: Optionally commit accepted suggestions using user’s token, then generate docs.
#     """
#     try:
#         if auth_token:
#             developer_agent.set_dynamic_token(auth_token)
#             print(f"🔐 Received dynamic GitHub token for user {user_id[:6]}...")

#         result = developer_agent.apply_changes_and_generate_docs(
#             repo_url=repo_url,
#             user_id=user_id,
#             accepted_suggestions=accepted_suggestions,
#             rejected_suggestions=rejected_suggestions,
#             doc_format=doc_format,
#             commit_changes=commit_changes,
#         )
#         return {"status": "success", "result": result}

#     except Exception as e:
#         print(f"❌ Error in apply-changes: {e}")
#         raise HTTPException(status_code=500, detail=str(e))


# # ========== EXPORT DOCUMENTATION ==========
# @app.post("/export-docs")
# async def export_docs(request: ExportRequest, preview: bool = False):
#     try:
#         print(f"\n[EXPORT] Exporting docs in format: {request.format.upper()}")
#         docs_output = [{"file_name": d.file_name, "documentation": d.documentation} for d in request.docs]
#         fmt = request.format.lower()
#         path = None

#         if fmt == "pdf":
#             from reportlab.pdfgen import canvas
#             from reportlab.lib.pagesizes import A4
#             path = tempfile.mktemp(suffix=".pdf")
#             c = canvas.Canvas(path, pagesize=A4)
#             width, height = A4
#             y = height - 50
#             for doc in docs_output:
#                 c.setFont("Helvetica-Bold", 14)
#                 c.drawString(50, y, f"{doc['file_name']}")
#                 y -= 20
#                 c.setFont("Helvetica", 11)
#                 for line in doc["documentation"].splitlines():
#                     if y < 50:
#                         c.showPage()
#                         y = height - 50
#                         c.setFont("Helvetica", 11)
#                     c.drawString(50, y, line[:110])
#                     y -= 15
#                 y -= 30
#             c.save()

#         elif fmt == "md":
#             path = tempfile.mktemp(suffix=".md")
#             with open(path, "w", encoding="utf-8") as f:
#                 for doc in docs_output:
#                     f.write(f"# {doc['file_name']}\n\n{doc['documentation']}\n\n---\n")

#         elif fmt == "txt":
#             path = tempfile.mktemp(suffix=".txt")
#             with open(path, "w", encoding="utf-8") as f:
#                 for doc in docs_output:
#                     f.write(f"{doc['file_name']}\n{'='*40}\n{doc['documentation']}\n\n")

#         elif fmt == "docx":
#             from docx import Document
#             path = tempfile.mktemp(suffix=".docx")
#             document = Document()
#             document.add_heading("CodeZen Documentation", level=1)
#             for doc in docs_output:
#                 document.add_heading(doc["file_name"], level=2)
#                 document.add_paragraph(doc["documentation"])
#                 document.add_page_break()
#             document.save(path)

#         elif fmt == "html":
#             path = tempfile.mktemp(suffix=".html")
#             html_content = (
#                 "<html><head><title>CodeZen Documentation</title></head><body>"
#                 "<h1>CodeZen Documentation</h1>"
#             )
#             for doc in docs_output:
#                 formatted = doc["documentation"].replace("\n", "<br>")
#                 html_content += f"<h2>{doc['file_name']}</h2><p>{formatted}</p><hr>"
#             html_content += "</body></html>"

#             with open(path, "w", encoding="utf-8") as f:
#                 f.write(html_content)

#             if preview:
#                 return {"status": "success", "html": html_content}

#         else:
#             raise HTTPException(status_code=400, detail="Invalid format. Use pdf, md, txt, docx, or html.")

#         filename = f"CodeZen_Docs.{fmt}"
#         media_map = {
#             "pdf": "application/pdf",
#             "md": "text/markdown",
#             "txt": "text/plain",
#             "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#             "html": "text/html",
#         }

#         print(f"✅ Export successful → {path}")
#         return FileResponse(path, media_type=media_map.get(fmt, "application/octet-stream"), filename=filename)

#     except Exception as e:
#         print(f"❌ Error exporting docs: {e}")
#         raise HTTPException(status_code=500, detail=str(e))


from fastapi import FastAPI, Request, Body, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, RedirectResponse, JSONResponse
from developer_agent import DeveloperAgent
from pydantic import BaseModel
from typing import Optional, List
import os, tempfile, requests
from dotenv import load_dotenv
load_dotenv()

app = FastAPI(title="CodeZen Backend", version="3.5")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # ⚠️ For production, restrict this to frontend domain
    allow_methods=["*"],
    allow_headers=["*"],
)

print("🚀 FastAPI Initialized")
developer_agent = DeveloperAgent()






# ========== MODELS ==========
class SuggestionResponse(BaseModel):
    file_name: str
    updated_content: Optional[str] = None
    accepted: bool
    comment: Optional[str] = None

class FileDoc(BaseModel):
    file_name: str
    documentation: str

class ExportRequest(BaseModel):
    format: str
    docs: List[FileDoc]


# ========== GITHUB OAUTH ==========
GITHUB_CLIENT_ID = os.getenv("GITHUB_CLIENT_ID")
GITHUB_CLIENT_SECRET = os.getenv("GITHUB_CLIENT_SECRET")
REDIRECT_URI = "http://localhost:8000/oauth/callback"
FRONTEND_URL = "http://localhost:5173"

@app.get("/login/github")
def github_login():
    github_auth_url = (
        "https://github.com/login/oauth/authorize"
        f"?client_id={GITHUB_CLIENT_ID}"
        f"&redirect_uri={REDIRECT_URI}"
        "&scope=repo,user:email"
        "&allow_signup=true"
    )
    print("🔗 Redirecting user to GitHub OAuth:", github_auth_url)
    return RedirectResponse(url=github_auth_url)


@app.get("/oauth/callback")
def github_callback(code: str):
    try:
        token_url = "https://github.com/login/oauth/access_token"
        payload = {
            "client_id": GITHUB_CLIENT_ID,
            "client_secret": GITHUB_CLIENT_SECRET,
            "code": code,
        }
        headers = {"Accept": "application/json"}
        response = requests.post(token_url, json=payload, headers=headers)
        token_data = response.json()

        if "access_token" not in token_data:
            raise Exception(token_data.get("error_description", "Failed to get token"))

        access_token = token_data["access_token"]
        print(f"✅ GitHub OAuth success — Token acquired for user: {access_token[:8]}...")
        return RedirectResponse(f"{FRONTEND_URL}?token={access_token}")

    except Exception as e:
        print(f"❌ OAuth callback error: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)


# ========== STAGE 1+2: ANALYSIS ==========
# @app.post("/run-analysis")
# async def run_analysis(repo_url: str = Body(..., embed=True)):
#     try:
#         result = developer_agent.run_full_pipeline(repo_url=repo_url, commit_changes=False)
#         return {"status": "success", "result": result}
#     except Exception as e:
#         print(f"❌ Error in run-analysis: {e}")
#         raise HTTPException(status_code=500, detail=str(e))
@app.post("/run-analysis")
async def run_analysis(repo_url: str = Body(..., embed=True)):
    try:
        result = developer_agent.run_full_pipeline(
            repo_url=repo_url,
            commit_changes=False,
            analysis_only=True
        )
        return {"status": "success", "result": result}
    except Exception as e:
        import traceback
        print("ERROR IN ANALYSIS")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))



# ========== STAGE 3+4: APPLY CHANGES + DOCS ==========

# @app.post("/apply-changes")
# async def apply_changes(
#     repo_url: str = Body(..., embed=True),
#     accepted_suggestions: Optional[List[SuggestionResponse]] = Body(default=None),
#     rejected_suggestions: Optional[List[SuggestionResponse]] = Body(default=None),
#     doc_format: str = Body("md", embed=True),
#     commit_changes: bool = Body(False, embed=True),
#     user_id: str = Body("default_user", embed=True),
#     auth_token: Optional[str] = Body(None, embed=True),
# ):
#     try:
#         if auth_token:
#             developer_agent.github_token = auth_token  # ✅ Dynamic assignment
#             print(f"🔐 Received dynamic GitHub token for user {user_id[:6]}...")

#         result = developer_agent.run_full_pipeline(
#             repo_url=repo_url,
#             user_id=user_id,
#             accepted_suggestions=accepted_suggestions,
#             rejected_suggestions=rejected_suggestions,
#             doc_format=doc_format,
#             commit_changes=commit_changes,
#         )
#         return {"status": "success", "result": result}

#     except Exception as e:
#         print(f"❌ Error in apply-changes: {e}")
#         raise HTTPException(status_code=500, detail=str(e))

@app.post("/apply-changes")
async def apply_changes(
    repo_url: str = Body(..., embed=True),
    accepted_suggestions: Optional[List[SuggestionResponse]] = Body(default=None),
    rejected_suggestions: Optional[List[SuggestionResponse]] = Body(default=None),
    doc_format: str = Body("md", embed=True),
    commit_changes: bool = Body(False, embed=True),
    user_id: str = Body("default_user", embed=True),
    auth_token: Optional[str] = Body(None, embed=True),
):
    try:
        if auth_token:
            developer_agent.set_dynamic_token(auth_token)
            print(f"🔐 Received dynamic GitHub token for user {user_id[:6]}...")

        # ✅ Convert Pydantic objects to plain dicts
        accepted = [s.dict() for s in accepted_suggestions or []]
        rejected = [s.dict() for s in rejected_suggestions or []]

        result = developer_agent.apply_changes_and_generate_docs(
            repo_url=repo_url,
            user_id=user_id,
            accepted_suggestions=accepted,
            rejected_suggestions=rejected,
            doc_format=doc_format,
            commit_changes=commit_changes,
            
        )
        return {"status": "success", "result": result}

    except Exception as e:
        print(f"❌ Error in apply-changes: {e}")
        raise HTTPException(status_code=500, detail=str(e))
 
from fastapi.responses import FileResponse

# @app.post("/generate-docs")
# async def generate_docs(request: Request):
#     data = await request.json()
#     repo_url = data.get("repo_url")
#     user_id = data.get("user_id", "default_user")
#     format_type = data.get("format", "md")  # pdf, docx, or md

#     try:
#         result = developer_agent.run_full_pipeline(
#             repo_url=repo_url,
#             user_id=user_id,
#             doc_format=format_type,
#             analysis_only=False,
#             use_cache=True
#         )

#         # 🔹 If the agent returned a file path (for PDF/DOCX), send it as a download
#         if isinstance(result.get("project_doc"), str) and os.path.exists(result["project_doc"]):
#             file_path = result["project_doc"]
#             filename = f"{result['repo']}_TDD.{format_type}"
#             media_map = {
#                 "pdf": "application/pdf",
#                 "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#                 "md": "text/markdown"
#             }
#             print(f"✅ [DOWNLOAD] Returning {format_type.upper()} file: {file_path}")
#             return FileResponse(file_path, media_type=media_map.get(format_type, "application/octet-stream"), filename=filename)


#         # 🔹 Otherwise (Markdown or string response)
#         return {"status": "success", "data": result}

#     except Exception as e:
#         print(f"❌ Error in generate-docs: {e}")
#         raise HTTPException(status_code=500, detail=str(e))
from starlette.background import BackgroundTask
@app.post("/generate-docs")
async def generate_docs(request: Request):
    data = await request.json()
    repo_url = data.get("repo_url")
    user_id = data.get("user_id", "default_user")
    format_type = data.get("format", "md")  # pdf, docx, or md

    print(f"\n⚙️ [generate-docs] Processing request for repo: {repo_url} | Format: {format_type.upper()}")

    try:
        result = developer_agent.run_full_pipeline(
            repo_url=repo_url,
            user_id=user_id,
            doc_format=format_type,
            analysis_only=False,
            use_cache=True
        )

        doc_content = result.get("project_doc")

        # --- 1. File Response Logic (PDF/DOCX) ---
        if isinstance(doc_content, str) and os.path.exists(doc_content):
            file_path = doc_content
            filename = f"{result.get('repo', 'CodeZen')}_TDD.{format_type}"
            media_map = {
                "pdf": "application/pdf",
                "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "md": "text/markdown"
            }
            print(f"✅ [DOWNLOAD] Returning {format_type.upper()} file: {file_path}")
            return FileResponse(
                file_path, 
                media_type=media_map.get(format_type, "application/octet-stream"), 
                filename=filename, 
                # Add a cleanup handler to delete the temp file after sending
                background=BackgroundTask(os.remove, file_path) 
            )

        # --- 2. Markdown/String Response Logic ---
        # Log a preview of the received documentation content
        preview = str(doc_content)[:100].replace('\n', ' ') if doc_content else 'None/Empty'
        print(f"✅ [DOC PREVIEW] First 100 chars of 'project_doc': {preview}")
        
        # Explicit check for empty Markdown content
        if format_type == "md" and (doc_content is None or doc_content.strip() == ""):
             print("❌ [DOCUMENTATION ERROR] Agent returned EMPTY 'project_doc' content.")
             # Raise an error to provide a clear status for the frontend
             raise HTTPException(
                 status_code=500, 
                 detail="The documentation agent ran but returned an empty document. Check agent logs."
             )

        # Return the nested JSON structure for Markdown (or other string data)
        return {"status": "success", "data": result}

    except HTTPException:
        # Re-raise explicit HTTP exceptions (like the 500 we raised above)
        raise
        
    except Exception as e:
        print(f"❌ Error in generate-docs: {e}")
        # Log the full traceback for better debugging
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")

# # ========== EXPORT DOCUMENTATION ==========
# @app.post("/export-docs")
# async def export_docs(request: ExportRequest, preview: bool = False):
#     try:
#         print(f"\n[EXPORT] Exporting docs in format: {request.format.upper()}")
#         docs_output = [{"file_name": d.file_name, "documentation": d.documentation} for d in request.docs]
#         fmt = request.format.lower()
#         path = None

#         if fmt == "pdf":
#             from reportlab.pdfgen import canvas
#             from reportlab.lib.pagesizes import A4
#             path = tempfile.mktemp(suffix=".pdf")
#             c = canvas.Canvas(path, pagesize=A4)
#             width, height = A4
#             y = height - 50
#             for doc in docs_output:
#                 c.setFont("Helvetica-Bold", 14)
#                 c.drawString(50, y, f"{doc['file_name']}")
#                 y -= 20
#                 c.setFont("Helvetica", 11)
#                 for line in doc["documentation"].splitlines():
#                     if y < 50:
#                         c.showPage()
#                         y = height - 50
#                         c.setFont("Helvetica", 11)
#                     c.drawString(50, y, line[:110])
#                     y -= 15
#                 y -= 30
#             c.save()

#         elif fmt == "md":
#             path = tempfile.mktemp(suffix=".md")
#             with open(path, "w", encoding="utf-8") as f:
#                 for doc in docs_output:
#                     f.write(f"# {doc['file_name']}\n\n{doc['documentation']}\n\n---\n")

#         elif fmt == "txt":
#             path = tempfile.mktemp(suffix=".txt")
#             with open(path, "w", encoding="utf-8") as f:
#                 for doc in docs_output:
#                     f.write(f"{doc['file_name']}\n{'='*40}\n{doc['documentation']}\n\n")

#         elif fmt == "docx":
#             from docx import Document
#             path = tempfile.mktemp(suffix=".docx")
#             document = Document()
#             document.add_heading("CodeZen Documentation", level=1)
#             for doc in docs_output:
#                 document.add_heading(doc["file_name"], level=2)
#                 document.add_paragraph(doc["documentation"])
#                 document.add_page_break()
#             document.save(path)

#         elif fmt == "html":
#             path = tempfile.mktemp(suffix=".html")
#             html_content = (
#                 "<html><head><title>CodeZen Documentation</title></head><body>"
#                 "<h1>CodeZen Documentation</h1>"
#             )
#             for doc in docs_output:
#                 formatted = doc["documentation"].replace("\n", "<br>")
#                 html_content += f"<h2>{doc['file_name']}</h2><p>{formatted}</p><hr>"
#             html_content += "</body></html>"

#             with open(path, "w", encoding="utf-8") as f:
#                 f.write(html_content)

#             if preview:
#                 return {"status": "success", "html": html_content}

#         else:
#             raise HTTPException(status_code=400, detail="Invalid format. Use pdf, md, txt, docx, or html.")

#         filename = f"CodeZen_Docs.{fmt}"
#         media_map = {
#             "pdf": "application/pdf",
#             "md": "text/markdown",
#             "txt": "text/plain",
#             "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#             "html": "text/html",
#         }

#         print(f"✅ Export successful → {path}")
#         return FileResponse(path, media_type=media_map.get(fmt, "application/octet-stream"), filename=filename)

#     except Exception as e:
#         print(f"❌ Error exporting docs: {e}")
#         raise HTTPException(status_code=500, detail=str(e))

@app.post("/export-docs")
async def export_docs(request: ExportRequest, preview: bool = False):
    from markdown import markdown  # ✅ Converts markdown -> HTML
    import tempfile
    from fastapi.responses import FileResponse
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from bs4 import BeautifulSoup  # ✅ To strip tags for PDF text
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.enum.style import WD_STYLE_TYPE

    try:
        print(f"\n[EXPORT] Exporting docs in format: {request.format.upper()}")
        fmt = request.format.lower()
        docs_output = [{"file_name": d.file_name, "documentation": d.documentation} for d in request.docs]

        # 🧠 Convert Markdown → HTML for all rich formats
        combined_html = ""
        for doc in docs_output:
            html_body = markdown(doc["documentation"])
            combined_html += f"<h2>{doc['file_name']}</h2>{html_body}<hr>"

        # --- Handle different formats ---
        path = None

        # ✅ PDF Export (formatted)
        if fmt == "pdf":
            from weasyprint import HTML
            path = tempfile.mktemp(suffix=".pdf")
            html_template = f"""
            <html>
            <head>
                <style>
                    body {{ font-family: 'Helvetica', sans-serif; margin: 40px; color: #111; }}
                    h1, h2 {{ color: #d32f2f; }}
                    hr {{ margin: 20px 0; }}
                    p {{ font-size: 12pt; line-height: 1.5; }}
                    ul {{ margin-left: 20px; }}
                </style>
            </head>
            <body>
                <h1>CodeZen Documentation</h1>
                {combined_html}
            </body>
            </html>
            """
            HTML(string=html_template).write_pdf(path)
       
        # ✅ DOCX Export (fully styled & polished)
        elif fmt == "docx":
        

            document = Document()

            # --- Cover Page ---
            document.add_paragraph().add_run("\n\n\n")
            cover = document.add_paragraph("CodeZen Technical Design Document")
            cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cover_run = cover.runs[0]
            cover_run.font.size = Pt(24)
            cover_run.font.bold = True
            cover_run.font.color.rgb = RGBColor(211, 47, 47)  # red

            document.add_paragraph().add_run("\nGenerated by CodeZen AI Developer Agent").italic = True
            document.add_page_break()

            # --- Global Styles ---
            style = document.styles["Normal"]
            font = style.font
            font.name = "Calibri"
            font.size = Pt(11)

            # --- Custom Styles ---
            for style_name, color, size, bold in [
                ("Heading1Custom", RGBColor(211, 47, 47), 18, True),
                ("Heading2Custom", RGBColor(233, 30, 99), 14, True),
                ("Heading3Custom", RGBColor(33, 150, 243), 12, True),
            ]:
                obj_style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                obj_style.font.name = "Calibri"
                obj_style.font.size = Pt(size)
                obj_style.font.bold = bold
                obj_style.font.color.rgb = color
                obj_style.paragraph_format.space_after = Pt(8)

            # --- Document Title ---
            document.add_paragraph("CodeZen Documentation", style="Heading1Custom")

            # --- Add Content ---
            for doc in docs_output:
                document.add_page_break()
                document.add_paragraph(doc["file_name"], style="Heading2Custom")

                soup = BeautifulSoup(markdown(doc["documentation"]), "html.parser")

                for element in soup.descendants:
                    if element.name == "h1":
                        document.add_paragraph(element.get_text(), style="Heading1Custom")
                    elif element.name == "h2":
                        document.add_paragraph(element.get_text(), style="Heading2Custom")
                    elif element.name == "h3":
                        document.add_paragraph(element.get_text(), style="Heading3Custom")
                    elif element.name == "p":
                        p = document.add_paragraph(element.get_text())
                        p.paragraph_format.space_after = Pt(6)
                    elif element.name == "ul":
                        for li in element.find_all("li", recursive=False):
                            document.add_paragraph(li.get_text(), style="List Bullet")
                    elif element.name == "ol":
                        for li in element.find_all("li", recursive=False):
                            document.add_paragraph(li.get_text(), style="List Number")
                    elif element.name == "code":
                        para = document.add_paragraph()
                        run = para.add_run(element.get_text())
                        run.font.name = "Consolas"
                        run.font.size = Pt(10)
                        shading_elm = run._element.rPr
                        run.font.color.rgb = RGBColor(255, 87, 34)
                    elif element.name == "hr":
                        document.add_paragraph("_" * 80)

            # --- Footer ---
            section = document.sections[-1]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = "Generated by CodeZen AI Developer Agent"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            path = tempfile.mktemp(suffix=".docx")
            document.save(path)

 

        # ✅ HTML Export (formatted)
        elif fmt == "html":
            html_template = f"""
            <html>
            <head>
                <title>CodeZen Documentation</title>
                <style>
                    body {{ font-family: 'Segoe UI', sans-serif; margin: 40px; background: #f8f8f8; color: #111; }}
                    h1, h2 {{ color: #c62828; }}
                    pre {{ background: #efefef; padding: 10px; border-radius: 8px; }}
                </style>
            </head>
            <body>
                <h1>CodeZen Documentation</h1>
                {combined_html}
            </body>
            </html>
            """
            path = tempfile.mktemp(suffix=".html")
            with open(path, "w", encoding="utf-8") as f:
                f.write(html_template)

        # ✅ Plain Text (no markdown symbols)
        elif fmt == "txt":
            path = tempfile.mktemp(suffix=".txt")
            with open(path, "w", encoding="utf-8") as f:
                for doc in docs_output:
                    soup = BeautifulSoup(markdown(doc["documentation"]), "html.parser")
                    text = soup.get_text()
                    f.write(f"{doc['file_name']}\n{'='*40}\n{text}\n\n")

        # ✅ Markdown (original format)
        elif fmt == "md":
            path = tempfile.mktemp(suffix=".md")
            with open(path, "w", encoding="utf-8") as f:
                for doc in docs_output:
                    f.write(f"# {doc['file_name']}\n\n{doc['documentation']}\n\n---\n")

        else:
            raise HTTPException(status_code=400, detail="Invalid format. Use pdf, docx, md, html, or txt.")

        filename = f"CodeZen_Docs.{fmt}"
        media_map = {
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "md": "text/markdown",
            "txt": "text/plain",
            "html": "text/html",
        }

        print(f"✅ Export successful → {path}")
        return FileResponse(path, media_type=media_map.get(fmt, "application/octet-stream"), filename=filename)

    except Exception as e:
        print(f"❌ Error exporting docs: {e}")
        raise HTTPException(status_code=500, detail=str(e))
